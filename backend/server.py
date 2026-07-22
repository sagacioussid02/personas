from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Request, BackgroundTasks
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import field_validator
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import copy
import os
import threading
import time
from collections import defaultdict, deque
from dotenv import load_dotenv
from typing import Optional, List, Dict, Any, Union
import json
import uuid
import re
from datetime import datetime
import boto3
from botocore.exceptions import ClientError
from pathlib import Path
from pypdf import PdfReader
import io
import asyncio
import hmac
import hashlib
import httpx
import jwt
from urllib.parse import quote
from agent_orchestrator import run_chat_orchestration
from personality_agent import detect_archetype, get_archetype, get_all_archetypes, review_response
from source_memory import (
    build_correction_source,
    build_deepen_sources,
    build_initial_sources,
    ensure_sources,
    merge_sources,
    record_knowledge_gap,
    topic_tags_for_question,
)

# Load environment variables
load_dotenv()


def _stable_source_id(source: Dict[str, Any]) -> str:
    source_type = str(source.get("source_type", "")).strip()
    content = str(source.get("content", "")).strip()
    title = str(source.get("title", "")).strip()
    url = str(source.get("url", "")).strip()
    raw_value = f"{source_type}\n{title}\n{url}\n{content}"
    return hashlib.sha256(raw_value.encode("utf-8")).hexdigest()


def _normalize_source_ids(sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized_sources = []
    for source in sources:
        normalized_source = dict(source)
        normalized_source["source_id"] = _stable_source_id(normalized_source)
        normalized_sources.append(normalized_source)
    return normalized_sources

app = FastAPI()

# Configure CORS
origins = os.getenv("CORS_ORIGINS", "http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=False,
    allow_methods=["GET", "POST", "PATCH", "OPTIONS"],
    allow_headers=["*"],
)

# Initialize Bedrock client - see Q42 on https://edwarddonner.com/faq if the Region gives you problems
bedrock_client = boto3.client(
    service_name="bedrock-runtime",
    region_name=os.getenv("DEFAULT_AWS_REGION", "us-east-1")
)

# Bedrock model selection - see Q42 on https://edwarddonner.com/faq for more
BEDROCK_MODEL_ID = os.getenv("BEDROCK_MODEL_ID", "global.amazon.nova-2-lite-v1:0")

# Memory storage configuration
USE_S3 = os.getenv("USE_S3", "false").lower() == "true"
PERSONALITY_REVIEW_ENABLED = os.getenv("PERSONALITY_REVIEW_ENABLED", "false").lower() == "true"
S3_BUCKET = os.getenv("S3_BUCKET", "")
MEMORY_DIR = os.getenv("MEMORY_DIR", "../memory")

# Initialize S3 client if needed
if USE_S3:
    if not S3_BUCKET:
        raise RuntimeError("USE_S3=true but S3_BUCKET environment variable is not set")
    s3_client = boto3.client("s3")

# Local twins dir: use /tmp/twins in Lambda (package dir is read-only), local path otherwise
_IN_LAMBDA = bool(os.getenv("AWS_LAMBDA_FUNCTION_NAME"))
TWINS_DIR = "/tmp/twins" if _IN_LAMBDA else os.path.join(os.path.dirname(__file__), "twins")  # nosec B108 — /tmp is the only writable path in Lambda; S3 is used when USE_S3=true
TWINS_S3_PREFIX = "twins/"

# Small per-recipient/per-twin marker objects for sharing and public-feature
# requests — same scoped-prefix pattern as TWINS_S3_PREFIX's per-user twin
# listing, so lookups stay O(1)-ish instead of scanning the whole bucket.
SHARES_DIR = "/tmp/shares" if _IN_LAMBDA else os.path.join(os.path.dirname(__file__), "shares")  # nosec B108
SHARES_S3_PREFIX = "shares/"  # shares/{email}/{twin_id}.json
PUBLIC_SHARE_REQUESTS_DIR = "/tmp/public_share_requests" if _IN_LAMBDA else os.path.join(os.path.dirname(__file__), "public_share_requests")  # nosec B108
PUBLIC_SHARE_REQUESTS_S3_PREFIX = "public_share_requests/"  # public_share_requests/{twin_id}.json
FEATURED_DIR = "/tmp/featured" if _IN_LAMBDA else os.path.join(os.path.dirname(__file__), "featured")  # nosec B108
FEATURED_S3_PREFIX = "featured/"  # featured/{twin_id}.json — marker that a twin is community-featured

_TWIN_ID_RE = re.compile(r'^[a-f0-9]{32}$')
_EMAIL_RE_FULL = re.compile(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$')

# ── Connect-to-creator notifications (AWS SES — no passwords, uses IAM role) ──
_SES_FROM_EMAIL = os.getenv("SES_FROM_EMAIL", "").strip()
_ADMIN_EMAILS = [e.strip() for e in os.getenv("ADMIN_EMAILS", "").split(",") if e.strip()]
_SES_REGION = (os.getenv("SES_REGION") or os.getenv("DEFAULT_AWS_REGION") or "us-east-1").strip()
_ses_client = boto3.client("ses", region_name=_SES_REGION) if _SES_FROM_EMAIL else None

_CONNECT_RE = re.compile(
    r'\b('
    r'(give|send|leave|share|have|pass)(?!\s+me\b)(\s+\w+){0,3}\s+feedback|'
    r'pass\s+(that|this|it|along)\s+(along|on)|'
    r'pass\s+along|'
    r'feedback\s+for\s+(you|your\s+team|this\s+app|the\s+site)|'
    r'contact\s+(you|us|sidd|the\s+creator|the\s+owner)|'
    r'reach\s+out|get\s+in\s+touch|'
    r'how\s+(do\s+i|can\s+i|to)\s+(contact|reach|connect)|'
    r'how\s+(do\s+i|can\s+i|to)\s+(send|give|share|leave)(\s+\w+){0,3}\s+feedback|'
    r'connect\s+with\s+(you|sidd|the\s+creator)|'
    r'talk\s+to\s+(you|sidd|the\s+real)'
    r')\b',
    re.IGNORECASE | re.DOTALL,
)
# Patterns to extract contact info shared voluntarily in chat by anonymous users
# Compiled with IGNORECASE so triggers match any capitalisation. False positives
# (e.g. "I'm not sure") are filtered in _extract_identity_from_history by
# requiring the first character of the captured name to be uppercase — words
# typed in title case almost always represent a proper name rather than a common
# word.
_NAME_IN_CHAT_RE = re.compile(
    r"(?i:my name is|call me|this is|i'?m|i am)\s+"
    r"([A-Z][a-z]{1,20}(?:\s+[A-Z][a-z]{1,20})?)"
)
_EMAIL_IN_CHAT_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")


def _extract_identity_from_history(
    history: list, current_message: str = ""
) -> tuple[Optional[str], Optional[str]]:
    """Scan user messages (oldest-first, plus current message) for a name/email.

    Returns (name, email) — either may be None if not found.
    Scans the current message first so a same-turn disclosure is captured.
    """
    candidates = [current_message] + [
        m.get("content", "") for m in history if m.get("role") == "user"
    ]
    name: Optional[str] = None
    email: Optional[str] = None
    for text in candidates:
        if not name:
            m = _NAME_IN_CHAT_RE.search(text)
            if m:
                captured = m.group(1).strip()
                # Require first char to be uppercase so common words after
                # "i'm" / "i am" (e.g. "not", "fine", "sure") are filtered out.
                if captured and captured[0].isupper():
                    name = captured.title()
        if not email:
            m = _EMAIL_IN_CHAT_RE.search(text)
            if m:
                email = m.group(0)
        if name and email:
            break
    return name, email


def _recent_transcript(conversation: list, current_message: str, max_messages: int = 8) -> str:
    """Render the last few turns as readable 'User: ... / Twin: ...' lines, plus
    the current triggering message. A connect/feedback notification quoting a
    single out-of-context line (e.g. "can I connect with you?") tells the twin
    owner nothing about what was actually discussed - this gives them the real
    conversation instead."""
    recent = conversation[-max_messages:]
    lines = []
    for msg in recent:
        role = "User" if msg.get("role") == "user" else "Twin"
        content = (msg.get("content") or "").strip()
        if content:
            lines.append(f"{role}: {content}")
    lines.append(f"User: {current_message}")
    return "\n".join(lines)


# ── Feedback notification rate limiting ───────────────────────────────────────
# Both anonymous and authenticated callers are capped on a rolling 7-day window
# to deter bot/DDoS abuse of the SES path. State lives in the warm Lambda
# container; cold starts reset it. For stricter enforcement across horizontal
# scale, move state to DynamoDB or Redis. SES daily quota is the hard ceiling.
def _get_int_env(name: str, default: int) -> int:
    value = os.getenv(name)
    if value is None:
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


_FEEDBACK_NOTIFY_RATE_LIMIT = _get_int_env("FEEDBACK_NOTIFY_RATE_LIMIT", 3)
_AUTH_FEEDBACK_NOTIFY_RATE_LIMIT = _get_int_env("AUTH_FEEDBACK_NOTIFY_RATE_LIMIT", 5)
_NOTIFY_WINDOW_SECONDS = 7 * 24 * 3600.0  # 7 days
_anon_notify_ip_history: dict[str, deque] = defaultdict(deque)
_anon_notify_session_seen: dict[str, float] = {}  # session_id -> timestamp; pruned on read
_anon_notify_session_email: dict[str, str] = {}  # session_id -> last email we already notified about
_auth_notify_user_history: dict[str, deque] = defaultdict(deque)
_anon_notify_lock = threading.Lock()
_auth_notify_lock = threading.Lock()

# Session-scoped "awaiting content" state: set when a message expresses
# connect/feedback intent with no substantive content yet (e.g. "I want to
# reach out to say something"), so the notification waits for the follow-up
# that actually says what — instead of firing immediately on the vague opener
# and burning the one-shot-per-session rate-limit slot before there's anything
# real to tell the twin owner. Short TTL since this should be answered in the
# same sitting; not the multi-day window the rate limiters use.
_PENDING_CONNECT_TTL_SECONDS = 30 * 60.0  # 30 minutes
_pending_connect_intent: dict[str, dict] = {}  # session_id -> {"message": str, "intent_type": str, "timestamp": float}
_pending_connect_lock = threading.Lock()


def _pop_pending_connect(session_id: str) -> Optional[dict]:
    """Return and clear this session's pending connect-intent entry, if any and not expired."""
    with _pending_connect_lock:
        now = time.time()
        expired = [sid for sid, entry in _pending_connect_intent.items() if now - entry["timestamp"] > _PENDING_CONNECT_TTL_SECONDS]
        for sid in expired:
            _pending_connect_intent.pop(sid, None)
        return _pending_connect_intent.pop(session_id, None)


def _set_pending_connect(session_id: str, message: str, intent_type: str) -> None:
    with _pending_connect_lock:
        _pending_connect_intent[session_id] = {
            "message": message,
            "intent_type": intent_type,
            "timestamp": time.time(),
        }


def _truncate_ip(ip: str) -> str:
    """Return a /24-truncated IPv4 address (last octet zeroed) for reduced PII.

    IPv6 addresses and non-parseable values are returned unchanged.
    Example: '1.2.3.4' → '1.2.3.x'
    """
    parts = ip.split(".")
    if len(parts) == 4 and all(p.isdigit() for p in parts):
        return f"{parts[0]}.{parts[1]}.{parts[2]}.x"
    return ip


def _client_ip(req: Request) -> str:
    # Only trust X-Forwarded-For when running behind API Gateway / Lambda (where
    # the header is injected by AWS and the raw socket is always a VPC hop).
    # In all other environments fall back to the direct socket address to prevent
    # clients from spoofing IPs to bypass the per-IP rate limiter.
    if _IN_LAMBDA:
        xff = req.headers.get("x-forwarded-for", "")
        if xff:
            return xff.split(",")[0].strip()
    return req.client.host if req.client else "unknown"


def _should_notify_anon(ip: str, session_id: str, email: Optional[str] = None) -> bool:
    """Return True iff an anonymous notification slot is available, consuming it.

    Limits: one notification per session UNLESS the message reveals a new email
    address we haven't already notified about for this session (e.g. the user's
    first message expresses vague interest in connecting, and a later message in
    the same session actually leaves contact info — that second one should still
    reach the inbox). A per-IP rolling 7-day cap of FEEDBACK_NOTIFY_RATE_LIMIT
    (default 3) is the hard ceiling in all cases, so repeatedly "discovering" new
    emails can't be used to bypass rate limiting. Returns False without consuming
    when the env flag is <= 0 or any cap is hit.
    """
    if _FEEDBACK_NOTIFY_RATE_LIMIT <= 0:
        return False
    with _anon_notify_lock:
        now = time.monotonic()
        # Prune expired session records to prevent unbounded growth
        expired = [sid for sid, ts in _anon_notify_session_seen.items() if now - ts > _NOTIFY_WINDOW_SECONDS]
        for sid in expired:
            _anon_notify_session_seen.pop(sid, None)
            _anon_notify_session_email.pop(sid, None)
        already_notified = session_id in _anon_notify_session_seen
        new_email = bool(email) and _anon_notify_session_email.get(session_id) != email
        if already_notified and not new_email:
            return False
        history = _anon_notify_ip_history[ip]
        while history and now - history[0] > _NOTIFY_WINDOW_SECONDS:
            history.popleft()
        if not history:
            # All entries expired; remove stale key to prevent unbounded dict growth
            _anon_notify_ip_history.pop(ip, None)
        if len(history) >= _FEEDBACK_NOTIFY_RATE_LIMIT:
            return False
        _anon_notify_ip_history[ip].append(now)  # re-creates via defaultdict if key was just removed
        _anon_notify_session_seen[session_id] = now
        if email:
            _anon_notify_session_email[session_id] = email
        return True


def _should_notify_auth(chatter_id: str) -> bool:
    """Return True iff an authenticated notification slot is available, consuming it.

    Limit: per-chatter_id rolling 7-day cap of
    AUTH_FEEDBACK_NOTIFY_RATE_LIMIT (default 5). Returns False without
    consuming when the env flag is <= 0 or the cap is hit.
    """
    if _AUTH_FEEDBACK_NOTIFY_RATE_LIMIT <= 0:
        return False
    with _auth_notify_lock:
        now = time.monotonic()
        history = _auth_notify_user_history[chatter_id]
        while history and now - history[0] > _NOTIFY_WINDOW_SECONDS:
            history.popleft()
        if not history:
            # All entries expired; remove stale key to prevent unbounded dict growth
            _auth_notify_user_history.pop(chatter_id, None)
        if len(history) >= _AUTH_FEEDBACK_NOTIFY_RATE_LIMIT:
            return False
        _auth_notify_user_history[chatter_id].append(now)  # re-creates via defaultdict if key was just removed
        return True


_NOTIFY_SUMMARY_MODEL_ID = "amazon.nova-micro-v1:0"
_NOTIFY_SUMMARY_PROMPT = (
    "In 1-3 sentences, in third person, summarize what this person wants to convey "
    "to the creator (Sidd) - their actual request, feedback, or message. Do not "
    "describe the conversation mechanics or repeat contact info (it's included "
    "separately). If there isn't much to summarize, say so plainly.\n\nConversation:\n"
)


async def _summarize_for_notification(transcript: str) -> str:
    """Generate a short, readable summary of the transcript for the notification
    email - replaces relying on the single-message intent classifier's optional
    one-liner, which is empty whenever the notification is resolved from a
    pending intent-only opener (that turn's classification was NO/INTENT, not
    CONTENT, so it never produced a summary)."""
    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=_NOTIFY_SUMMARY_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": _NOTIFY_SUMMARY_PROMPT + transcript}]}],
            inferenceConfig={"maxTokens": 120, "temperature": 0},
        )
        return response["output"]["message"]["content"][0]["text"].strip()
    except Exception as exc:
        print(f"[notify] Summary generation failed: {exc}")
        return ""


async def _notify_connect_intent(
    transcript: str,
    identity_block: str,
    session_id: str,
    twin_name: str,
    intent_type: str = "connect",
) -> None:
    """Fire-and-forget SES email when a user asks to connect with the creator."""
    if not _ses_client or not _ADMIN_EMAILS:
        return
    try:
        type_label = {
            "feedback": "Feedback",
            "compliment": "Compliment",
            "complaint": "Complaint",
            "review": "Review request",
            "connect": "Connect request",
        }.get(intent_type, "Connect request")
        summary = await _summarize_for_notification(transcript)
        summary_line = f"Summary: {summary}\n" if summary else ""
        body = (
            f"{type_label} received via {twin_name}\n\n"
            f"{summary_line}"
            f"Persona: {twin_name}\n"
            f"Session: {session_id}\n"
            f"Time: {datetime.utcnow().isoformat()}Z\n\n"
            f"Contact info:\n{identity_block}\n\n"
            f"Full conversation:\n{transcript}\n"
        )
        await asyncio.to_thread(
            _ses_client.send_email,
            Source=_SES_FROM_EMAIL,
            Destination={"ToAddresses": _ADMIN_EMAILS},
            Message={
                "Subject": {"Data": f"[Personas] {type_label} via {twin_name}"},
                "Body": {"Text": {"Data": body}},
            },
        )
        print(f"[notify] Connect alert sent for session {session_id} (type={intent_type})")
    except Exception as exc:
        print(f"[notify] Failed to send connect alert: {exc}")


# Intent classification model — cheapest/fastest
_INTENT_MODEL_ID = "amazon.nova-micro-v1:0"
_INTENT_PROMPT = (
    "Classify the following message about the user's intent toward the creator. "
    "Reply in exactly one of these formats:\n"
    "INTENT|<type>|<one-sentence summary>\n"
    "  — the user wants to give feedback, leave a review, share a compliment, "
    "lodge a complaint, or connect/contact the creator, but THIS message does not "
    "itself contain the actual content to relay (e.g. \"can I connect with you?\", "
    "\"I want to reach out to say something\", \"how do I contact you?\").\n"
    "CONTENT|<type>|<one-sentence summary>\n"
    "  — THIS message itself contains the substantive point, feedback, or request "
    "the user wants passed along (e.g. \"tell Sidd I'm impressed by his work\", "
    "\"I loved the AI safety talk\", a description of a job opportunity).\n"
    "NO\n"
    "  — none of the above.\n"
    "<type> is one of: feedback, compliment, complaint, review, connect\n\nMessage: "
)


async def _classify_feedback_intent(message: str) -> tuple[bool, str, str, bool]:
    """
    Returns (is_connect, intent_type, summary, has_content).
    has_content distinguishes "wants to connect, hasn't said what about yet"
    (INTENT) from "this message is the actual thing to relay" (CONTENT) - see
    _decide_connect_notification, which only ever notifies once has_content
    is True for some message in the exchange.
    Uses a Bedrock Nova Micro classifier (temperature=0, maxTokens=40).
    Falls back to _CONNECT_RE (treated as CONTENT, preserving prior behavior)
    if the Bedrock call fails.
    """
    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=_INTENT_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": _INTENT_PROMPT + message}]}],
            inferenceConfig={"maxTokens": 40, "temperature": 0},
        )
        answer = response["output"]["message"]["content"][0]["text"].strip()
        upper = answer.upper()
        if upper.startswith("INTENT") or upper.startswith("CONTENT"):
            has_content = upper.startswith("CONTENT")
            parts = answer.split("|", 2)
            intent_type = parts[1].strip().lower() if len(parts) > 1 else "connect"
            summary = parts[2].strip() if len(parts) > 2 else ""
            print(f"[intent] LLM classified message as {'CONTENT' if has_content else 'INTENT'} (type={intent_type})")
            return True, intent_type, summary, has_content
        print("[intent] LLM classified message as NORMAL")
        return False, "", "", False
    except Exception as exc:
        print(f"[intent] LLM classification failed, falling back to regex: {exc}")
        is_connect = bool(_CONNECT_RE.search(message))
        return is_connect, "connect" if is_connect else "", "", is_connect


def _decide_connect_notification(
    message: str,
    session_id: str,
    viewer_is_authenticated: bool,
    chatter_id: Optional[str],
    user_email: Optional[str],
    ip: str,
    conversation: list,
    is_connect: bool,
    has_content: bool,
    intent_type: str,
) -> tuple[bool, Optional[str], Optional[str], str, str]:
    """
    Synchronously decide whether this message earns a real SES notification, and
    consume the rate-limit slot if so. Pure in-memory bookkeeping — no I/O — so
    it's safe to call before building the user-facing response, which means the
    acknowledgment we show can honestly reflect whether we're actually sending
    something instead of always claiming success.

    Returns (should_notify, transcript, identity_block, status, effective_intent_type).
    transcript/identity_block are None when should_notify is False.
    status is one of:
      "none"       — not connect-related, and no earlier turn is pending either;
                     no acknowledgment should be shown.
      "awaiting"   — intent expressed but this message carries no content yet;
                     pending state recorded, no acknowledgment should be shown
                     (the twin's own reply already asks what they'd like to share).
      "notified"   — real content present (either this message, or the very next
                     message resolving an earlier pending intent-only opener);
                     should_notify reflects whether the rate limit actually
                     allowed sending it.
      "suppressed" — real content present but rate-limited; no email sent.
    effective_intent_type is intent_type when content arrived this turn, or the
    pending entry's type when resolving a previously-deferred intent-only opener.

    Once a pending entry exists, the very next user message always resolves it —
    the user was just asked "what would you like to share?", so whatever they
    say next is treated as the answer regardless of how this message classifies
    on its own. Relying on the classifier to independently re-recognize the
    follow-up as CONTENT is what let a real answer ("...add more public
    personas... I am Alex") get re-classified as another content-less INTENT,
    silently deferring the notification an extra turn.
    """
    pending = _pop_pending_connect(session_id)

    if pending:
        effective_intent_type = pending["intent_type"]
    elif is_connect and not has_content:
        # Wants to connect/give feedback, but hasn't said what about yet — don't
        # notify (and don't consume the rate-limit slot) on a vague opener like
        # "I want to reach out to say something". Wait for the very next message,
        # which resolves unconditionally above on the next call.
        _set_pending_connect(session_id, message, intent_type)
        return False, None, None, "awaiting", intent_type
    elif is_connect and has_content:
        effective_intent_type = intent_type
    else:
        return False, None, None, "none", ""

    if viewer_is_authenticated:
        should_notify = _should_notify_auth(chatter_id)
        if not should_notify:
            print(f"[notify] Auth notification suppressed by rate limit (chatter_id={chatter_id})")
            return False, None, None, "suppressed", effective_intent_type
        identity = f"Email: {user_email}" if user_email else f"chatter_id: {chatter_id}"
        transcript = _recent_transcript(conversation, message)
        return True, transcript, identity, "notified", effective_intent_type

    name, shared_email = _extract_identity_from_history(conversation, message)
    should_notify = _should_notify_anon(ip, session_id, email=shared_email)
    if not should_notify:
        print(f"[notify] Anon notification suppressed by rate limit (ip={ip}, session={session_id})")
        return False, None, None, "suppressed", effective_intent_type
    identity_lines = []
    if name:
        identity_lines.append(f"Name (from chat): {name}")
    identity_lines.append(
        f"Email (from chat): {shared_email}" if shared_email
        else "Email: not shared yet — the twin should ask for this if the conversation continues"
    )
    identity_lines.append(f"IP: {_truncate_ip(ip)}")
    identity_lines.append(f"Session: {session_id}")
    transcript = _recent_transcript(conversation, message)
    return True, transcript, "\n".join(identity_lines), "notified", effective_intent_type


async def _send_connect_notification(
    transcript: str,
    identity_block: str,
    session_id: str,
    twin_name: str,
    intent_type: str = "connect",
) -> None:
    """Background task: actually send the SES alert (including the summary
    generation, itself a Bedrock call). Rate limiting was already decided (and
    the slot consumed) synchronously by _decide_connect_notification before
    this was scheduled, so nothing here affects the user-facing response."""
    try:
        await asyncio.wait_for(
            _notify_connect_intent(transcript, identity_block, session_id, twin_name, intent_type),
            timeout=8.0,
        )
    except Exception:
        pass

# ── Public personas ────────────────────────────────────────────────────────────
# Loaded once at startup from backend/public_personas/*.json.
# These are served to anyone (no auth) and use stable hard-coded twin_ids.
_PUBLIC_PERSONAS_DIR = os.path.join(os.path.dirname(__file__), "public_personas")
_PUBLIC_PERSONAS: dict[str, dict] = {}  # keyed by twin_id

def _load_public_personas() -> None:
    directory = Path(_PUBLIC_PERSONAS_DIR)
    if not directory.exists():
        print(
            f"Warning: public personas directory '{_PUBLIC_PERSONAS_DIR}' not found; "
            "the /public-personas endpoint will be empty. "
            "If running in AWS Lambda, ensure 'public_personas/' is included in the deployment package."
        )
        return
    if not directory.is_dir():
        print(
            f"Warning: public personas path '{_PUBLIC_PERSONAS_DIR}' exists but is not a directory; "
            "skipping public persona loading."
        )
        return
    for f in directory.glob("*.json"):
        try:
            data = json.loads(f.read_text())
            if not isinstance(data, dict):
                print(f"Warning: public persona file {f.name} does not contain a JSON object; skipping")
                continue
            tid = data.get("twin_id")
            name = data.get("name")
            if not isinstance(tid, str) or not _TWIN_ID_RE.match(tid):
                print(f"Warning: public persona file {f.name} has invalid or missing twin_id; skipping")
                continue
            if not isinstance(name, str) or not name.strip():
                print(f"Warning: public persona file {f.name} has invalid or missing name; skipping")
                continue
            data.setdefault("is_public", True)
            data.setdefault("user_id", None)
            _PUBLIC_PERSONAS[tid] = data
        except Exception as exc:
            print(f"Warning: could not load public persona {f.name}: {exc}")

_load_public_personas()

# Max questions an anonymous user may ask a public persona before being prompted to sign up
PUBLIC_PERSONA_ANON_LIMIT = 5

# Secret used to derive opaque session keys — must be set in production.
# Generate with: python -c "import secrets; print(secrets.token_hex(32))"
SESSION_HMAC_SECRET = os.getenv("SESSION_HMAC_SECRET", "")
if not SESSION_HMAC_SECRET and (_IN_LAMBDA or os.getenv("USE_S3", "").lower() in ("1", "true") or os.getenv("ENVIRONMENT", "").lower() == "prod"):
    raise RuntimeError(
        "SESSION_HMAC_SECRET environment variable must be set in production. "
        "Generate one with: python -c \"import secrets; print(secrets.token_hex(32))\""
    )

# --- Clerk JWT auth ---
CLERK_JWKS_URL = os.getenv("CLERK_JWKS_URL", "")
# Derive issuer from JWKS URL: strip /.well-known/jwks.json
CLERK_ISSUER = CLERK_JWKS_URL.removesuffix("/.well-known/jwks.json") if CLERK_JWKS_URL else ""
# Optional: set CLERK_AUDIENCE if your Clerk app has a custom audience configured
CLERK_AUDIENCE = os.getenv("CLERK_AUDIENCE", "") or None
_jwks_cache: Optional[dict] = None
_bearer = HTTPBearer(auto_error=False)


async def _fetch_jwks() -> dict:
    if not CLERK_JWKS_URL:
        raise HTTPException(status_code=500, detail="Auth not configured")
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(CLERK_JWKS_URL)
            resp.raise_for_status()
    except httpx.TimeoutException as exc:
        raise HTTPException(status_code=503, detail="JWKS endpoint timeout") from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=503, detail="JWKS endpoint unavailable") from exc
    try:
        return resp.json()
    except ValueError as exc:
        raise HTTPException(status_code=503, detail="Invalid JWKS response") from exc


async def _get_jwks(force_refresh: bool = False) -> dict:
    global _jwks_cache
    if not force_refresh and _jwks_cache:
        return _jwks_cache
    _jwks_cache = await _fetch_jwks()
    return _jwks_cache


def _find_key(jwks: dict, kid: str) -> Optional[dict]:
    keys = jwks.get("keys")
    if not isinstance(keys, list):
        raise HTTPException(status_code=503, detail="JWKS payload invalid")
    return next((k for k in keys if k.get("kid") == kid), None)


async def _decode_user_claims(
    credentials: Optional[HTTPAuthorizationCredentials],
) -> tuple[Optional[str], Optional[str]]:
    """Decode user_id and email from credentials.

    Returns (user_id, email) tuple — either field may be None on missing
    credentials or JWT validation failures (expired, bad signature, unknown kid).
    Propagates 5xx HTTPExceptions from auth infrastructure (JWKS fetch timeout,
    misconfigured CLERK_JWKS_URL, etc.) so callers can surface them correctly
    rather than masking outages as 401s.
    """
    if not credentials:
        return None, None
    token = credentials.credentials
    try:
        header = jwt.get_unverified_header(token)
        kid = header.get("kid", "")

        jwks = await _get_jwks()
        key = _find_key(jwks, kid)

        if key is None:
            jwks = await _get_jwks(force_refresh=True)
            key = _find_key(jwks, kid)
        if key is None:
            return None, None

        public_key = jwt.algorithms.RSAAlgorithm.from_jwk(json.dumps(key))
        decode_options: dict = {}
        if not CLERK_AUDIENCE:
            decode_options["verify_aud"] = False
        payload = jwt.decode(
            token,
            public_key,
            algorithms=["RS256"],
            issuer=CLERK_ISSUER if CLERK_ISSUER else None,
            audience=CLERK_AUDIENCE,
            options=decode_options,
        )
        user_id: str = payload.get("sub", "")
        email: str = payload.get("email", "")
        return user_id or None, email or None
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        # Token is bad/expired — treat as anonymous, don't raise
        return None, None
    except HTTPException as exc:
        if exc.status_code >= 500:
            raise  # Auth infrastructure failure — propagate so it surfaces correctly
        return None, None  # 4xx from auth layer — treat as invalid token


async def _decode_user_id(
    credentials: Optional[HTTPAuthorizationCredentials],
) -> Optional[str]:
    """Convenience wrapper — returns just the user_id from _decode_user_claims."""
    user_id, _ = await _decode_user_claims(credentials)
    return user_id


async def get_current_user_id(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(_bearer),
) -> str:
    """Strict auth — raises 401 if token is missing or invalid. Use for protected endpoints.
    5xx auth infrastructure errors (JWKS outage, misconfiguration) propagate as-is.
    """
    if not credentials:
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_id = await _decode_user_id(credentials)
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")
    return user_id


async def get_current_user_email(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(_bearer),
) -> str:
    """Strict auth — raises 401 if token is missing/invalid, 400 if the token's
    Clerk account has no email claim (shouldn't happen in practice, but the
    sharing feature depends on it)."""
    if not credentials:
        raise HTTPException(status_code=401, detail="Not authenticated")
    _user_id, email = await _decode_user_claims(credentials)
    if not _user_id:
        raise HTTPException(status_code=401, detail="Invalid token")
    if not email:
        raise HTTPException(status_code=400, detail="Account has no email on file")
    return _normalize_email(email)


async def get_current_admin_email(
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(_bearer),
) -> str:
    """Admin-only endpoints: reuses ADMIN_EMAILS (already the SES notification
    destination list) as the authorization source, rather than introducing a
    separate role system for what is, today, a single-admin app."""
    email = await get_current_user_email(credentials)
    if email not in {e.lower() for e in _ADMIN_EMAILS}:
        raise HTTPException(status_code=403, detail="Admin access required")
    return email

# Expected keys for the personality model returned by /create-twin
_PERSONALITY_MODEL_KEYS = {
    "core_values", "decision_heuristics", "risk_profile",
    "what_they_optimize_for", "what_they_avoid",
    "communication_traits", "blind_spots",
    "decision_framework", "personality_summary",
}


def _extract_json_object(text: str, required_key: str | None = None) -> dict:
    """Extract the first complete JSON object from *text*.

    Scans every '{' position in turn so that a stray JSON fragment appended
    after natural-language text doesn't shadow the real response object.
    Returns the first valid dict (optionally containing *required_key*).
    """
    pos = 0
    last_error: Exception = ValueError("No JSON object found in response")
    while True:
        start = text.find('{', pos)
        if start == -1:
            raise last_error
        depth, in_string, escape_next = 0, False, False
        end = None
        for i, ch in enumerate(text[start:], start):
            if escape_next:
                escape_next = False
                continue
            if ch == '\\' and in_string:
                escape_next = True
                continue
            if ch == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    end = i
                    break
        if end is None:
            # Record the error but continue scanning from the next position
            last_error = ValueError("Unbalanced braces — could not extract JSON object")
            pos = start + 1
            continue
        try:
            candidate = json.loads(text[start:end + 1])
            if isinstance(candidate, dict):
                if required_key is None or required_key in candidate:
                    return candidate
        except json.JSONDecodeError as exc:
            last_error = exc
        pos = end + 1  # advance past this block and try the next '{'


def _extract_json_array(text: str) -> list:
    """Extract the first complete JSON array using balanced-bracket scan."""
    start = text.find('[')
    if start == -1:
        raise ValueError("No JSON array found in response")
    depth, in_string, escape_next = 0, False, False
    for i, ch in enumerate(text[start:], start):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '[':
            depth += 1
        elif ch == ']':
            depth -= 1
            if depth == 0:
                return json.loads(text[start:i + 1])
    raise ValueError("Unbalanced brackets — could not extract JSON array")


def _s3_get_twin(key: str) -> Optional[dict]:
    """Fetch and parse a twin JSON from S3 by key. Returns None on missing key."""
    try:
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=key)
        return json.loads(response["Body"].read().decode("utf-8"))
    except ClientError as e:
        if e.response["Error"]["Code"] == "NoSuchKey":
            return None
        raise


def load_twin(twin_id: str) -> Optional[dict]:
    """Load a saved twin's data by ID. Validates ID format and confines path to TWINS_DIR.

    Checks public personas first (in-memory), then S3 / local disk.
    S3 layout: flat key twins/{twin_id}.json for O(1) public lookup.
    Per-user key twins/{user_id}/{twin_id}.json exists in parallel for listing.
    """
    if not _TWIN_ID_RE.match(twin_id):
        raise HTTPException(status_code=400, detail="Invalid twin ID format")

    if twin_id in _PUBLIC_PERSONAS:
        return copy.deepcopy(_PUBLIC_PERSONAS[twin_id])

    if USE_S3:
        # Direct flat-key lookup — O(1), safe for public endpoints
        return _s3_get_twin(f"{TWINS_S3_PREFIX}{twin_id}.json")

    path = os.path.realpath(os.path.join(TWINS_DIR, f"{twin_id}.json"))
    if not path.startswith(os.path.realpath(TWINS_DIR) + os.sep):
        raise HTTPException(status_code=400, detail="Invalid twin ID")
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return None


# Request/Response models
class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    twin_id: Optional[str] = None  # if set, chat with a user-created twin


class ChatSource(BaseModel):
    source_id: str
    source_type: str
    title: str
    snippet: str
    confidence: str
    tags: List[str] = Field(default_factory=list)
    matched_terms: List[str] = Field(default_factory=list)


class ChatGrounding(BaseModel):
    answer_type: str
    confidence_label: str
    grounding_mode: str


class ChatResponse(BaseModel):
    response: str
    session_id: str
    grounding: Optional[ChatGrounding] = None
    sources: List[ChatSource] = Field(default_factory=list)


class Message(BaseModel):
    role: str
    content: str
    timestamp: str


class CreateTwinRequest(BaseModel):
    name: str
    title: str
    bio: str
    email: str = ""
    skills: str = ""
    experience: str = ""
    achievements: str = ""
    coreValues: str = ""
    decisionStyle: str = ""
    riskTolerance: str = ""
    pastDecisions: str = ""
    communicationStyle: str = ""
    writingSamples: str = ""
    blindSpots: str = ""
    archetype_id: Optional[str] = None
    responseStyle: Optional[str] = "balanced"
    verbalQuirks: Optional[str] = ""
    linkedinParsed: Optional[Dict[str, Any]] = None

    @field_validator("name", "title", "bio")
    @classmethod
    def strip_and_require(cls, v: str, info) -> str:
        v = v.strip()
        if not v:
            raise ValueError(f"{info.field_name} must not be empty")
        limits = {"name": 100, "title": 150, "bio": 2000}
        limit = limits.get(info.field_name, 2000)
        if len(v) > limit:
            raise ValueError(f"{info.field_name} must be {limit} characters or fewer")
        return v

    @field_validator(
        "email", "skills", "experience", "achievements",
        "coreValues", "decisionStyle", "riskTolerance", "pastDecisions",
        "communicationStyle", "writingSamples", "blindSpots",
    )
    @classmethod
    def strip_optional(cls, v: str, info) -> str:
        v = v.strip()
        limits = {
            "email": 254,       # RFC 5321 max
            "skills": 1000,
            "experience": 5000,
            "achievements": 2000,
            "coreValues": 2000,
            "decisionStyle": 3000,
            "riskTolerance": 500,
            "pastDecisions": 3000,
            "communicationStyle": 2000,
            "writingSamples": 1000,
            "blindSpots": 2000,
        }
        limit = limits.get(info.field_name, 2000)
        if len(v) > limit:
            raise ValueError(f"{info.field_name} must be {limit} characters or fewer")
        return v

    @field_validator("archetype_id")
    @classmethod
    def strip_archetype_id(cls, v: Optional[str]) -> Optional[str]:
        if v is None:
            return v
        v = v.strip()
        if len(v) > 50:
            raise ValueError("archetype_id must be 50 characters or fewer")
        return v or None

    @field_validator("responseStyle")
    @classmethod
    def strip_response_style(cls, v: Optional[str]) -> Optional[str]:
        if not v:
            return "balanced"
        v_normalized = v.strip().lower()
        allowed_styles = {"concise", "balanced", "detailed"}
        if v_normalized not in allowed_styles:
            return "balanced"
        return v_normalized

    @field_validator("verbalQuirks")
    @classmethod
    def strip_verbal_quirks(cls, v: Optional[str]) -> Optional[str]:
        if v is None:
            return ""
        v = v.strip()
        if len(v) > 1500:
            raise ValueError("verbalQuirks must be 1500 characters or fewer")
        return v


# Valid session_id shapes:
#   - UUID (any version) from str(uuid.uuid4()): xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx
#     Note: the regex accepts any lowercase hex UUID, not only v4.
#   - 64-char hex from HMAC-SHA256 / SHA-256 _derive_session_id
_SESSION_ID_RE = re.compile(r'^(?:[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}|[0-9a-f]{64})$')


def _validate_session_id(session_id: str) -> None:
    """Raise 400 if session_id doesn't match the allowed format, preventing path traversal."""
    if not _SESSION_ID_RE.match(session_id):
        raise HTTPException(status_code=400, detail="Invalid session ID format")


# Memory management functions
def get_memory_path(session_id: str) -> str:
    return f"sessions/{session_id}.json"


def _load_raw_record(session_id: str) -> Optional[Union[List[Dict], Dict[str, Any]]]:
    """Load the raw conversation record from storage (S3 or local).

    Returns the parsed JSON value (list or dict) if found, or None if the session
    does not exist. Raises on storage errors (e.g. AccessDenied, throttling).
    """
    _validate_session_id(session_id)
    if USE_S3:
        try:
            response = s3_client.get_object(Bucket=S3_BUCKET, Key=get_memory_path(session_id))
            return json.loads(response["Body"].read().decode("utf-8"))
        except ClientError as e:
            if e.response["Error"]["Code"] == "NoSuchKey":
                # Fallback: check legacy flat key (pre-sessions/ prefix)
                try:
                    response = s3_client.get_object(Bucket=S3_BUCKET, Key=f"{session_id}.json")
                    return json.loads(response["Body"].read().decode("utf-8"))
                except ClientError as legacy_e:
                    if legacy_e.response["Error"]["Code"] == "NoSuchKey":
                        return None
                    raise  # AccessDenied, throttling, etc. — surface, don't silently drop
            else:
                raise
    else:
        sessions_dir = os.path.realpath(os.path.join(MEMORY_DIR, "sessions"))
        file_path = os.path.realpath(os.path.join(sessions_dir, f"{session_id}.json"))
        if not file_path.startswith(sessions_dir + os.sep):
            raise HTTPException(status_code=400, detail="Invalid session ID format")
        if os.path.exists(file_path):
            with open(file_path, "r") as f:
                return json.load(f)
        # Fallback: check legacy flat path (pre-sessions/ prefix)
        memory_dir_real = os.path.realpath(MEMORY_DIR)
        legacy_path = os.path.realpath(os.path.join(MEMORY_DIR, f"{session_id}.json"))
        if not legacy_path.startswith(memory_dir_real + os.sep):
            raise HTTPException(status_code=400, detail="Invalid session ID format")
        if os.path.exists(legacy_path):
            with open(legacy_path, "r") as f:
                return json.load(f)
        return None


def load_conversation(session_id: str) -> List[Dict]:
    """Load conversation history from storage. Handles legacy list format and current dict format."""
    raw = _load_raw_record(session_id)
    if raw is None:
        return []
    # Support both legacy list format and current dict format
    if isinstance(raw, list):
        return raw
    return raw.get("messages", [])


def save_conversation(
    session_id: str,
    messages: List[Dict],
    chatter_id: Optional[str] = None,
    twin_owner_id: Optional[str] = None,
):
    """Save conversation history with owner metadata for future access-control migration."""
    _validate_session_id(session_id)
    data: Any = {
        "session_id": session_id,
        "chatter_id": chatter_id,       # authenticated user who is chatting (None = anonymous)
        "twin_owner_id": twin_owner_id,  # user_id of the twin's creator
        "messages": messages,
    }
    if USE_S3:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=get_memory_path(session_id),
            Body=json.dumps(data, indent=2),
            ContentType="application/json",
        )
    else:
        sessions_dir = os.path.realpath(os.path.join(MEMORY_DIR, "sessions"))
        os.makedirs(sessions_dir, exist_ok=True)
        file_path = os.path.realpath(os.path.join(sessions_dir, f"{session_id}.json"))
        if not file_path.startswith(sessions_dir + os.sep):
            raise HTTPException(status_code=400, detail="Invalid session ID format")
        with open(file_path, "w") as f:
            json.dump(data, f, indent=2)


@app.get("/")
async def root():
    return {
        "message": "AI Digital Twin API (Powered by AWS Bedrock)",
        "memory_enabled": True,
        "storage": "S3" if USE_S3 else "local",
        "ai_model": BEDROCK_MODEL_ID
    }


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "use_s3": USE_S3,
        "bedrock_model": BEDROCK_MODEL_ID
    }


@app.get("/archetypes")
async def list_archetypes():
    """Return all available archetypes for the frontend dropdown."""
    return {"archetypes": get_all_archetypes()}


def _derive_session_id(chatter_id: str, twin_id: str) -> str:
    """Return an opaque, stable session key for an authenticated user + twin pair.

    HMAC-SHA256 of 'chatter_id:twin_id' with SESSION_HMAC_SECRET makes the key
    non-guessable even if both IDs are known. This keeps session identifiers
    opaque for authenticated callers (e.g. when fetching /conversation/{session_id})
    and avoids leaking information about underlying user or twin IDs.
    Falls back to SHA-256 without a secret when SESSION_HMAC_SECRET is not set
    (local development only) — preserves stability and format validity but not secrecy.
    """
    if SESSION_HMAC_SECRET:
        return hmac.new(
            SESSION_HMAC_SECRET.encode("utf-8"),
            f"{chatter_id}:{twin_id}".encode("utf-8"),
            hashlib.sha256,
        ).hexdigest()
    # Dev fallback — no secret, but still produces a valid 64-char hex so the
    # session_id passes format validation. Not safe for production (predictable).
    return hashlib.sha256(f"{chatter_id}:{twin_id}".encode("utf-8")).hexdigest()


@app.post("/chat", response_model=ChatResponse)
async def chat(
    request: ChatRequest,
    http_request: Request,
    background_tasks: BackgroundTasks,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(_bearer),
):
    try:
        # Resolve caller identity (optional — anonymous callers are allowed)
        chatter_id, user_email = await _decode_user_claims(credentials)

        # Derive an opaque stable session key for authenticated users so memory
        # persists across devices and page reloads. Anonymous users fall back to
        # the client-supplied session_id (ephemeral, within-browser-session only).
        if chatter_id and request.twin_id:
            session_id = _derive_session_id(chatter_id, request.twin_id)
        else:
            session_id = request.session_id or str(uuid.uuid4())

        # Guard against session hijacking: if the stored record was created by
        # an authenticated user, only that same user (or a caller with no stored
        # chatter_id to compare against) may continue it.
        existing_record = _load_raw_record(session_id)
        stored_chatter_id = (
            existing_record.get("chatter_id")
            if isinstance(existing_record, dict)
            else None
        )
        if stored_chatter_id is not None and stored_chatter_id != chatter_id:
            raise HTTPException(
                status_code=403,
                detail="Forbidden: this session belongs to a different user",
            )

        # Load conversation history from the already-fetched record (avoids a
        # second storage read).
        if existing_record is None:
            conversation: List[Dict] = []
        elif isinstance(existing_record, list):
            conversation = existing_record
        else:
            conversation = existing_record.get("messages", [])

        # Load twin personality model if twin_id provided
        personality_model = None
        twin_name = None
        twin_title = None
        twin_data = None
        if request.twin_id:
            twin_data = load_twin(request.twin_id)
            # Enforce anonymous question limit on public personas
            if twin_data and twin_data.get("is_public") and not chatter_id:
                # Require a client-supplied session_id so the limit can't be
                # bypassed by omitting it (which would otherwise cause the
                # server to generate a fresh UUID, resetting the counter).
                if not request.session_id:
                    raise HTTPException(
                        status_code=400,
                        detail="session_id is required for anonymous public persona chat",
                    )
                anon_q_count = sum(1 for m in conversation if m.get("role") == "user")
                if anon_q_count >= PUBLIC_PERSONA_ANON_LIMIT:
                    raise HTTPException(
                        status_code=402,
                        detail="ANON_LIMIT_REACHED",
                    )
            if not twin_data:
                raise HTTPException(status_code=404, detail=f"Twin '{request.twin_id}' not found")
            personality_model = twin_data.get("personality_model", {})
            # Attach raw fields so context builder can access them
            personality_model["_context"] = twin_data.get("personality_model", {}).get("_context", {})
            twin_name = twin_data.get("name")
            twin_title = twin_data.get("title")

        # Determine response_style from personality model context.
        # Default twin (no twin_id) uses "concise" to keep homepage chat snappy.
        response_style = "concise" if not request.twin_id else "balanced"
        if personality_model:
            response_style = personality_model.get("_context", {}).get("responseStyle", response_style)

        corrections = twin_data.get("corrections") if twin_data else None
        normalized_sources = None
        if twin_data:
            normalized_sources = _normalize_source_ids(ensure_sources(twin_data))
            twin_data["sources"] = normalized_sources

        viewer_is_authenticated = chatter_id is not None

        # Run intent classification concurrently with chat orchestration.
        # Nova Micro (~300 ms) finishes well before Nova Lite (~2 s), so there
        # is no perceptible latency increase.
        # intent_summary is unused now — the notification email generates its own
        # summary from the full transcript at send time (see _summarize_for_notification),
        # since the per-message classifier's summary is empty whenever notification
        # is resolved from a pending intent-only opener rather than this exact turn.
        (is_connect, intent_type, _intent_summary, has_content), orchestration = await asyncio.gather(
            _classify_feedback_intent(request.message),
            asyncio.to_thread(
                run_chat_orchestration,
                twin_data=twin_data,
                sources=normalized_sources,
                user_message=request.message,
                conversation=conversation,
                bedrock_client=bedrock_client,
                model_id=BEDROCK_MODEL_ID,
                personality_model=personality_model,
                twin_name=twin_name,
                twin_title=twin_title,
                response_style=response_style,
                corrections=corrections,
                viewer_is_authenticated=viewer_is_authenticated,
            ),
        )
        assistant_response = orchestration["answer"]

        # Decide — synchronously, no I/O — whether this message actually earns a
        # notification, before we say anything to the user about it. This is what
        # makes the acknowledgment below honest instead of always claiming success.
        should_notify, notify_transcript, notify_identity, notify_status, effective_intent_type = _decide_connect_notification(
            message=request.message,
            session_id=session_id,
            viewer_is_authenticated=viewer_is_authenticated,
            chatter_id=chatter_id,
            user_email=user_email,
            ip=_client_ip(http_request),
            conversation=list(conversation),
            is_connect=is_connect,
            has_content=has_content,
            intent_type=intent_type,
        )

        # Acknowledge in the response only once there's something real to say —
        # "awaiting" (intent expressed, no content yet) gets no system-appended
        # acknowledgment; the twin's own reply already asks what they'd like to
        # share (see context.py's critical rules), so adding one here would
        # prematurely claim Sidd's been notified before there's anything to notify.
        if notify_status == "notified":
            ack_map = {
                "feedback": "I've passed your feedback along to Sidd — he'll appreciate hearing it.",
                "compliment": "I've shared your kind words with Sidd — that really means a lot.",
                "complaint": "I've forwarded your concern to Sidd so he can look into it.",
                "review": "I've let Sidd know you'd like to leave a review — thanks for taking the time.",
                "connect": "I've let Sidd know you'd like to get in touch — he'll reach out.",
            }
            ack = ack_map.get(effective_intent_type, ack_map["connect"])
            assistant_response = f"{assistant_response}\n\n*{ack}*"
        elif notify_status == "suppressed":
            ack = (
                "I've already flagged this conversation to Sidd, so I won't send a "
                "duplicate alert — it's saved here and he'll see it when he checks in."
            )
            assistant_response = f"{assistant_response}\n\n*{ack}*"

        # The rate-limit decision (and slot consumption) already happened above;
        # this background task only performs the actual SES network call.
        if should_notify:
            background_tasks.add_task(
                _send_connect_notification,
                notify_transcript,
                notify_identity,
                session_id,
                twin_name or "Sidd",
                effective_intent_type,
            )

        # Knowledge-gap ledger: record topics this twin answered weakly, so a
        # future deepen session can target what's actually worth asking about
        # instead of a fixed generic questionnaire. Pure, twin-agnostic check
        # here; the actual read-modify-write is deferred to a background task.
        if twin_data and request.twin_id and _is_gap_worthy(orchestration):
            background_tasks.add_task(
                _record_chat_gap,
                request.twin_id,
                topic_tags_for_question(request.message),
                request.message[:200],
            )

        # Personality review step (gated — enable via PERSONALITY_REVIEW_ENABLED=true)
        if PERSONALITY_REVIEW_ENABLED:
            archetype_id = twin_data.get("archetype_id") if request.twin_id and twin_data else None
            archetype = get_archetype(archetype_id) if archetype_id else None
            if archetype:
                twin_context = f"{twin_name or 'Professional'}, {twin_title or ''}. {twin_data.get('personality_model', {}).get('personality_summary', '')[:200]}"
                assistant_response = review_response(assistant_response, archetype, twin_context, bedrock_client, BEDROCK_MODEL_ID)

        # Update conversation history
        conversation.append(
            {"role": "user", "content": request.message, "timestamp": datetime.now().isoformat()}
        )
        conversation.append(
            {
                "role": "assistant",
                "content": assistant_response,
                "timestamp": datetime.now().isoformat(),
            }
        )

        # Save conversation — preserve an existing non-null chatter_id so that
        # unauthenticated retries cannot "de-own" a session that was previously
        # created by an authenticated caller.
        twin_owner_id = twin_data.get("user_id") if twin_data else None
        effective_chatter_id = chatter_id or stored_chatter_id
        save_conversation(session_id, conversation, chatter_id=effective_chatter_id, twin_owner_id=twin_owner_id)

        # Grounding/confidence badges and source snippets are an internal
        # transparency aid for whoever manages the twin, not end-user-facing
        # chat UI - a visitor (signed in or anonymous) shouldn't see raw
        # "medium confidence / grounded / Based on Profile summary" scaffolding
        # while talking to what should read as the person themselves. Visible
        # only to the twin's authenticated owner; nobody sees it for twins
        # with no owner configured (public personas, the default twin before
        # DEFAULT_TWIN_OWNER_USER_ID is set).
        can_view_source_details = (
            chatter_id is not None and twin_owner_id is not None and chatter_id == twin_owner_id
        )

        return ChatResponse(
            response=assistant_response,
            session_id=session_id,
            grounding=(
                ChatGrounding(**orchestration["grounding"])
                if can_view_source_details and orchestration["grounding"]
                else None
            ),
            sources=(
                [ChatSource(**source) for source in orchestration["retrieved_sources"]]
                if can_view_source_details
                else []
            ),
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in chat endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")


@app.get("/conversation/{session_id}")
async def get_conversation(
    session_id: str,
    _user_id: str = Depends(get_current_user_id),
):
    """Retrieve conversation history for the given session_id.

    Requires authentication. Enforces ownership: only the authenticated user
    whose ``chatter_id`` matches the stored record may retrieve the conversation.
    Returns 404 for unknown sessions, sessions without ownership metadata, or
    sessions owned by a different user.
    """
    try:
        raw = _load_raw_record(session_id)
        if raw is None:
            raise HTTPException(status_code=404, detail="Conversation not found")

        # Enforce ownership using stored chatter_id metadata.
        if isinstance(raw, dict) and "chatter_id" in raw:
            if raw["chatter_id"] != _user_id:
                # Hide existence details from unauthorized callers.
                raise HTTPException(status_code=404, detail="Conversation not found")
            messages = raw.get("messages", [])
        else:
            # Legacy list format or missing chatter_id — deny to avoid leaking data
            # from conversations not clearly associated with this user.
            raise HTTPException(status_code=404, detail="Conversation not found")

        return {"session_id": session_id, "messages": messages}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Cache for taglines to reduce API calls
tagline_cache = {"taglines": None, "timestamp": None}
TAGLINE_CACHE_TTL = 3600  # Cache for 1 hour


@app.get("/taglines")
async def get_taglines():
    """Generate humorous and attractive taglines dynamically using AI (cached)"""
    global tagline_cache
    from datetime import datetime, timedelta

    # Check if cache is still valid
    if (tagline_cache["taglines"] is not None and
        tagline_cache["timestamp"] is not None and
        datetime.now() - tagline_cache["timestamp"] < timedelta(seconds=TAGLINE_CACHE_TTL)):
        print("Returning cached taglines")
        return {"taglines": tagline_cache["taglines"]}

    try:
        print("Generating new taglines from AI...")
        prompt_text = """Generate exactly 10 short, humorous, witty, and attractive taglines for an AI Digital Twin product.
        Each tagline should be catchy, fun, and appeal to tech-savvy users. They should relate to AI, digital clones, productivity, or self-improvement.
        Keep each tagline to 2-5 words maximum.
        Format: Return only a JSON array of strings, nothing else.
        Example format: ["Coffee with this guy", "Resumes are old school", "Your digital brainpower unleashed"]"""

        messages = [
            {
                "role": "user",
                "content": [{"text": prompt_text}]
            }
        ]

        response = bedrock_client.converse(
            modelId=BEDROCK_MODEL_ID,
            messages=messages,
            inferenceConfig={
                "maxTokens": 500,
                "temperature": 0.9,
                "topP": 0.95
            }
        )

        response_text = response["output"]["message"]["content"][0]["text"]

        try:
            taglines = _extract_json_array(response_text)
            tagline_cache["taglines"] = taglines
            tagline_cache["timestamp"] = datetime.now()
            return {"taglines": taglines}
        except (ValueError, json.JSONDecodeError):
            fallback = ["Talk to your AI twin", "Your digital self awaits", "AI collaboration unlocked"]
            tagline_cache["taglines"] = fallback
            tagline_cache["timestamp"] = datetime.now()
            return {"taglines": fallback}

    except Exception as e:
        print(f"Error generating taglines: {str(e)}")
        # Return fallback taglines on error
        fallback = [
            "Coffee with this guy",
            "Resumes are old school",
            "Your digital brainpower unleashed",
            "The future of collaboration is here",
            "AI that gets you",
            "Your second brain in action",
            "Talk to your smarter self",
            "Meet Sidd 2.0",
            "Intelligence amplified",
            "Your AI just leveled up"
        ]
        tagline_cache["taglines"] = fallback
        tagline_cache["timestamp"] = datetime.now()
        return {"taglines": fallback}


@app.post("/parse-linkedin")
async def parse_linkedin(file: UploadFile = File(...)):
    """Extract structured profile data from a LinkedIn PDF using Bedrock"""
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    try:
        contents = await file.read()
        reader = PdfReader(io.BytesIO(contents))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from PDF")

    extract_prompt = f"""You are extracting structured profile data from a LinkedIn PDF export.

LinkedIn PDF content:
{text[:6000]}

Extract the following fields and return ONLY a valid JSON object with these exact keys:
{{
  "name": "full name",
  "title": "current job title and company",
  "bio": "2-3 sentence professional summary",
  "skills": "comma-separated list of top skills",
  "experience": "bullet list of roles: Company (dates): what they did",
  "achievements": "notable achievements, awards, or highlights",
  "communicationStyle": "inferred communication style based on their writing and background"
}}

If a field cannot be determined, use an empty string. Return only the JSON, no other text."""

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": extract_prompt}]}],
            inferenceConfig={"maxTokens": 1500, "temperature": 0.2},
        )
        response_text = response["output"]["message"]["content"][0]["text"]

        try:
            parsed = _extract_json_object(response_text)
        except (ValueError, json.JSONDecodeError) as exc:
            print(f"JSON extraction failed in parse-linkedin: {exc}")
            raise HTTPException(status_code=500, detail="Could not parse AI response as JSON")

        # Detect archetype from title
        archetype = detect_archetype(parsed.get("title", ""))
        parsed["archetype_id"] = archetype["id"] if archetype else None
        parsed["archetype_display_name"] = archetype["display_name"] if archetype else None

        return parsed

    except HTTPException:
        raise
    except ClientError as e:
        print(f"Bedrock ClientError in parse-linkedin: {e}")
        raise HTTPException(status_code=500, detail="Failed to process LinkedIn profile")
    except Exception as e:
        print(f"Unexpected error in parse-linkedin: {e}")
        raise HTTPException(status_code=500, detail="Failed to process LinkedIn profile")


@app.get("/twin/{twin_id}")
async def get_twin(twin_id: str):
    """Fetch public profile data for a twin"""
    twin_data = load_twin(twin_id)
    if not twin_data:
        raise HTTPException(status_code=404, detail="Twin not found")
    return {
        "twin_id": twin_data["twin_id"],
        "name": twin_data["name"],
        "title": twin_data.get("title", ""),
        "personality_summary": twin_data.get("personality_model", {}).get("personality_summary", ""),
        "core_values": twin_data.get("personality_model", {}).get("core_values", []),
        "archetype_display_name": twin_data.get("archetype_display_name"),
        "created_at": twin_data.get("created_at", ""),
        "source_count": len(ensure_sources(twin_data)),
    }


def _public_persona_summary(p: dict) -> dict:
    return {
        "twin_id": p["twin_id"],
        "persona_id": p.get("persona_id"),
        "name": p["name"],
        "title": p.get("title", ""),
        "tagline": p.get("tagline", ""),
        "era": p.get("era", ""),
        "image_url": p.get("image_url"),
        "personality_summary": p.get("personality_model", {}).get("personality_summary", ""),
        "chat_url": f"/twin?id={p['twin_id']}",
    }


@app.get("/public-personas")
async def list_public_personas():
    """Return the built-in public personas plus any user-created twin that has
    been approved for public featuring (see /twin/{id}/request-public and
    /admin/public-personas/{id}/approve)."""
    built_in = [
        _public_persona_summary(p)
        for p in sorted(
            _PUBLIC_PERSONAS.values(),
            key=lambda persona: (
                str(persona.get("persona_id") or ""),
                str(persona.get("name") or ""),
            ),
        )
    ]
    featured_markers = _small_object_list(FEATURED_S3_PREFIX, FEATURED_DIR)
    featured = []
    for marker in featured_markers:
        twin_data = load_twin(marker.get("twin_id", ""))
        if twin_data and twin_data.get("public_share_status") == "approved":
            summary = _public_persona_summary(twin_data)
            summary["persona_id"] = summary["persona_id"] or twin_data["twin_id"]
            featured.append(summary)
    return {"personas": built_in + featured}


@app.get("/users/me/twins")
async def list_my_twins(user_id: str = Depends(get_current_user_id)):
    """List all twins belonging to the authenticated user."""
    twins = []
    if USE_S3:
        # Scoped prefix — only fetches this user's objects, not a full table scan
        user_prefix = f"{TWINS_S3_PREFIX}{user_id}/"
        paginator = s3_client.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=user_prefix):
            for obj in page.get("Contents", []):
                try:
                    resp = s3_client.get_object(Bucket=S3_BUCKET, Key=obj["Key"])
                    data = json.loads(resp["Body"].read())
                    # Safety guard — prefix already scopes to this user
                    if data.get("user_id") == user_id:
                        twins.append({
                            "twin_id": data["twin_id"],
                            "name": data["name"],
                            "title": data.get("title", ""),
                            "archetype_display_name": data.get("archetype_display_name"),
                            "created_at": data.get("created_at", ""),
                            "chat_url": data.get("chat_url", f"/twin?id={data['twin_id']}"),
                            "depth_score": _compute_depth_score(data),
                        })
                except Exception as e:
                    print(f"Warning: could not read S3 object {obj['Key']}: {e}")
                    continue
    else:
        twins_path = Path(TWINS_DIR)
        if twins_path.exists():
            for f in sorted(twins_path.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
                try:
                    data = json.loads(f.read_text())
                    if data.get("user_id") == user_id:
                        twins.append({
                            "twin_id": data["twin_id"],
                            "name": data["name"],
                            "title": data.get("title", ""),
                            "archetype_display_name": data.get("archetype_display_name"),
                            "created_at": data.get("created_at", ""),
                            "chat_url": data.get("chat_url", f"/twin?id={data['twin_id']}"),
                            "depth_score": _compute_depth_score(data),
                        })
                except Exception as e:
                    print(f"Warning: could not read twin file {f}: {e}")
                    continue
    twins.sort(key=lambda t: t["created_at"], reverse=True)
    return {"twins": twins}


def _twin_summary(twin_data: dict) -> dict:
    """Shared shape for a twin summary row — used by list_my_twins-adjacent
    endpoints (shared-twins listing, share record snapshots)."""
    return {
        "twin_id": twin_data["twin_id"],
        "name": twin_data["name"],
        "title": twin_data.get("title", ""),
        "archetype_display_name": twin_data.get("archetype_display_name"),
        "created_at": twin_data.get("created_at", ""),
        "chat_url": twin_data.get("chat_url", f"/twin?id={twin_data['twin_id']}"),
    }


@app.delete("/twin/{twin_id}")
async def delete_twin(twin_id: str, user_id: str = Depends(get_current_user_id)):
    """Delete a twin the caller owns. Also revokes any outstanding shares and
    withdraws a pending/approved public-share so no dangling references to a
    now-deleted twin remain in the shares/public_share_requests/featured indexes."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    for email in twin_data.get("shared_with", []):
        _small_object_delete(SHARES_S3_PREFIX, SHARES_DIR, f"{email}/{twin_id}.json")
    _small_object_delete(PUBLIC_SHARE_REQUESTS_S3_PREFIX, PUBLIC_SHARE_REQUESTS_DIR, f"{twin_id}.json")
    _small_object_delete(FEATURED_S3_PREFIX, FEATURED_DIR, f"{twin_id}.json")

    _delete_twin_storage(twin_id, user_id)
    return {"status": "deleted", "twin_id": twin_id}


@app.get("/users/me/shared-twins")
async def list_shared_twins(email: str = Depends(get_current_user_email)):
    """List twins that have been shared with the caller's account email."""
    records = _small_object_list(SHARES_S3_PREFIX, SHARES_DIR, sub_prefix=f"{email}/")
    records.sort(key=lambda r: r.get("shared_at", ""), reverse=True)
    return {"twins": records}


class ShareTwinRequest(BaseModel):
    email: str

    @field_validator("email")
    @classmethod
    def email_valid(cls, v: str) -> str:
        v = v.strip()
        if not _EMAIL_RE_FULL.match(v):
            raise ValueError("Invalid email address")
        return v


@app.post("/twin/{twin_id}/share")
async def share_twin(
    twin_id: str,
    request: ShareTwinRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Share a twin the caller owns with another user by email. That user
    will see it under /users/me/shared-twins once they sign in — no
    invitation/acceptance step, matching the "just works when they log in"
    behavior asked for."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    email = _normalize_email(request.email)
    shared_with = list(twin_data.get("shared_with", []))
    if email not in shared_with:
        shared_with.append(email)
        twin_data["shared_with"] = shared_with
        _save_twin(twin_id, user_id, twin_data)

    record = _twin_summary(twin_data)
    record["shared_at"] = datetime.now().isoformat()
    record["shared_by"] = user_id
    _small_object_put(SHARES_S3_PREFIX, SHARES_DIR, f"{email}/{twin_id}.json", record)
    return {"status": "shared", "twin_id": twin_id, "email": email, "shared_with": shared_with}


@app.delete("/twin/{twin_id}/share")
async def unshare_twin(
    twin_id: str,
    request: ShareTwinRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Revoke a twin share."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    email = _normalize_email(request.email)
    shared_with = [e for e in twin_data.get("shared_with", []) if e != email]
    twin_data["shared_with"] = shared_with
    _save_twin(twin_id, user_id, twin_data)
    _small_object_delete(SHARES_S3_PREFIX, SHARES_DIR, f"{email}/{twin_id}.json")
    return {"status": "unshared", "twin_id": twin_id, "email": email, "shared_with": shared_with}


@app.post("/twin/{twin_id}/request-public")
async def request_public_feature(twin_id: str, user_id: str = Depends(get_current_user_id)):
    """Request that a twin the caller owns be featured on the public homepage.
    Takes effect only once an admin approves it via /admin/public-personas/{id}/approve -
    this just files the request and notifies the admin."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")
    if twin_data.get("public_share_status") == "pending":
        return {"status": "pending", "twin_id": twin_id}
    if twin_data.get("public_share_status") == "approved":
        return {"status": "approved", "twin_id": twin_id}

    twin_data["public_share_status"] = "pending"
    _save_twin(twin_id, user_id, twin_data)

    record = _twin_summary(twin_data)
    record["requested_by"] = user_id
    record["requested_at"] = datetime.now().isoformat()
    _small_object_put(PUBLIC_SHARE_REQUESTS_S3_PREFIX, PUBLIC_SHARE_REQUESTS_DIR, f"{twin_id}.json", record)

    if _ses_client and _ADMIN_EMAILS:
        try:
            await asyncio.wait_for(
                asyncio.to_thread(
                    _ses_client.send_email,
                    Source=_SES_FROM_EMAIL,
                    Destination={"ToAddresses": _ADMIN_EMAILS},
                    Message={
                        "Subject": {"Data": f"[Personas] Public-feature request: {twin_data.get('name')}"},
                        "Body": {"Text": {"Data": (
                            f"{twin_data.get('name')} ({twin_data.get('title', '')}) has been requested "
                            f"for public featuring.\n\nTwin ID: {twin_id}\nRequested by user: {user_id}\n\n"
                            f"Review at /admin.\n"
                        )}},
                    },
                ),
                timeout=5.0,
            )
        except Exception as exc:
            print(f"[notify] Failed to send public-feature request alert: {exc}")

    return {"status": "pending", "twin_id": twin_id}


@app.get("/admin/pending-public-personas")
async def list_pending_public_personas(_admin_email: str = Depends(get_current_admin_email)):
    records = _small_object_list(PUBLIC_SHARE_REQUESTS_S3_PREFIX, PUBLIC_SHARE_REQUESTS_DIR)
    records.sort(key=lambda r: r.get("requested_at", ""))
    return {"requests": records}


@app.post("/admin/public-personas/{twin_id}/approve")
async def approve_public_persona(twin_id: str, _admin_email: str = Depends(get_current_admin_email)):
    twin_data = load_twin(twin_id)
    if not twin_data:
        raise HTTPException(status_code=404, detail="Twin not found")

    twin_data["public_share_status"] = "approved"
    twin_data["is_public"] = True
    owner_id = twin_data.get("user_id")
    if owner_id:
        _save_twin(twin_id, owner_id, twin_data)
    else:
        _save_twin_flat(twin_id, twin_data)

    _small_object_delete(PUBLIC_SHARE_REQUESTS_S3_PREFIX, PUBLIC_SHARE_REQUESTS_DIR, f"{twin_id}.json")
    _small_object_put(FEATURED_S3_PREFIX, FEATURED_DIR, f"{twin_id}.json", {"twin_id": twin_id, "approved_at": datetime.now().isoformat()})
    return {"status": "approved", "twin_id": twin_id}


@app.post("/admin/public-personas/{twin_id}/reject")
async def reject_public_persona(twin_id: str, _admin_email: str = Depends(get_current_admin_email)):
    twin_data = load_twin(twin_id)
    if not twin_data:
        raise HTTPException(status_code=404, detail="Twin not found")

    twin_data["public_share_status"] = "rejected"
    owner_id = twin_data.get("user_id")
    if owner_id:
        _save_twin(twin_id, owner_id, twin_data)
    else:
        _save_twin_flat(twin_id, twin_data)

    _small_object_delete(PUBLIC_SHARE_REQUESTS_S3_PREFIX, PUBLIC_SHARE_REQUESTS_DIR, f"{twin_id}.json")
    return {"status": "rejected", "twin_id": twin_id}


def synthesize_personality_model(
    *,
    name: str,
    title: str,
    bio: str,
    skills: str = "",
    experience: str = "",
    achievements: str = "",
    coreValues: str = "",
    decisionStyle: str = "",
    riskTolerance: str = "",
    pastDecisions: str = "",
    communicationStyle: str = "",
    writingSamples: str = "",
    blindSpots: str = "",
) -> Dict[str, Any]:
    """Call Bedrock to synthesize a personality_model from profile fields.

    Shared by /create-twin and the default-twin migration script so both
    produce structurally identical models from the same prompt. Raises
    ValueError on any failure (bad Bedrock response, missing keys); callers
    translate that into their own error handling (HTTPException for the
    endpoint, a plain failure for the script).
    """
    synthesis_prompt = f"""You are building a personality model for an AI twin. Your job is to deeply analyze everything provided and produce a structured JSON model that captures how this person THINKS and DECIDES — not just what they've done.

This model will be used to answer questions like "What would {name} do?" in real situations.

=== PROFILE DATA ===

Name: {name}
Title: {title}
Bio: {bio}

Skills: {skills}

Work Experience:
{experience}

Achievements:
{achievements}

Core Values:
{coreValues}

Decision-Making Style:
{decisionStyle}

Risk Tolerance: {riskTolerance}

Past Decisions & Reasoning:
{pastDecisions}

Communication Style:
{communicationStyle}

Writing Samples/Links:
{writingSamples}

Blind Spots & Biases:
{blindSpots}

=== TASK ===

Analyze all of the above and return ONLY a valid JSON object with this exact structure:

{{
  "core_values": ["value 1", "value 2", ...],
  "decision_heuristics": [
    "When facing X type of decision, they tend to Y",
    ...
  ],
  "risk_profile": "one paragraph describing how they approach risk and uncertainty",
  "what_they_optimize_for": ["thing 1", "thing 2", ...],
  "what_they_avoid": ["thing 1", "thing 2", ...],
  "communication_traits": ["trait 1", "trait 2", ...],
  "blind_spots": ["blind spot 1", "blind spot 2", ...],
  "decision_framework": "2-3 sentence summary of their overall decision-making philosophy",
  "personality_summary": "3-4 sentence paragraph capturing the essence of who this person is and how they operate — written in second person as if talking to their twin"
}}

Be specific and concrete. Avoid generic statements. Infer from the data even when not explicit. Return only the JSON."""

    try:
        response = bedrock_client.converse(
            modelId=BEDROCK_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": synthesis_prompt}]}],
            inferenceConfig={"maxTokens": 2000, "temperature": 0.3},
        )
    except ClientError as e:
        raise ValueError(f"Bedrock error: {str(e)}")

    response_text = response["output"]["message"]["content"][0]["text"]
    try:
        personality_model = _extract_json_object(response_text)
    except (ValueError, json.JSONDecodeError):
        raise ValueError("Could not parse personality model from AI response")

    missing = _PERSONALITY_MODEL_KEYS - personality_model.keys()
    if missing:
        raise ValueError(f"Personality model missing expected keys: {missing}")

    return personality_model


@app.post("/create-twin")
async def create_twin(request: CreateTwinRequest, user_id: str = Depends(get_current_user_id)):
    """Synthesize submitted profile data into a structured personality model via Bedrock"""

    try:
        personality_model = synthesize_personality_model(
            name=request.name,
            title=request.title,
            bio=request.bio,
            skills=request.skills,
            experience=request.experience,
            achievements=request.achievements,
            coreValues=request.coreValues,
            decisionStyle=request.decisionStyle,
            riskTolerance=request.riskTolerance,
            pastDecisions=request.pastDecisions,
            communicationStyle=request.communicationStyle,
            writingSamples=request.writingSamples,
            blindSpots=request.blindSpots,
        )
    except ValueError as e:
        raise HTTPException(status_code=500, detail=str(e))

    twin_id = uuid.uuid4().hex  # 32 hex chars (128-bit) — no truncation, no collision risk

    # Embed the context fields the prompt builder needs directly in the personality model.
    personality_model["_context"] = {
        "bio": request.bio,
        "skills": request.skills,
        "experience": request.experience,
        "achievements": request.achievements,
        "coreValues": request.coreValues,
        "decisionStyle": request.decisionStyle,
        "pastDecisions": request.pastDecisions,
        "communicationStyle": request.communicationStyle,
        "blindSpots": request.blindSpots,
        "verbalQuirks": request.verbalQuirks or "",
        "responseStyle": request.responseStyle or "balanced",
    }

    # Resolve archetype — reject unknown IDs so clients aren't misled
    archetype_id = request.archetype_id or None
    archetype_obj = get_archetype(archetype_id) if archetype_id else None
    if archetype_id and archetype_obj is None:
        raise HTTPException(status_code=400, detail=f"Unknown archetype_id: {archetype_id!r}")
    archetype_display_name = archetype_obj["display_name"] if archetype_obj else None

    twin_data: Dict[str, Any] = {
        "twin_id": twin_id,
        "user_id": user_id,
        "name": request.name,
        "title": request.title,
        "archetype_id": archetype_id,
        "archetype_display_name": archetype_display_name,
        "personality_model": personality_model,
        "sources": build_initial_sources(request.model_dump(), request.linkedinParsed),
        "created_at": datetime.now().isoformat(),
        "chat_url": f"/twin?id={twin_id}",
    }

    if USE_S3:
        payload = json.dumps(twin_data, indent=2)
        # Flat key for O(1) public lookup (load_twin, /twin/{id}, /chat)
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=f"{TWINS_S3_PREFIX}{twin_id}.json",
            Body=payload,
            ContentType="application/json",
        )
        # Per-user key for efficient user listing (list_my_twins)
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=f"{TWINS_S3_PREFIX}{user_id}/{twin_id}.json",
            Body=payload,
            ContentType="application/json",
        )
    else:
        # Local / Lambda /tmp fallback — not durable across Lambda cold starts
        os.makedirs(TWINS_DIR, exist_ok=True)
        try:
            max_twins = int(os.getenv("MAX_TWINS_FILES", "1000"))
        except ValueError:
            max_twins = 1000
        existing = sorted(Path(TWINS_DIR).glob("*.json"), key=lambda p: p.stat().st_mtime)
        for old_file in existing[: max(0, len(existing) - max_twins + 1)]:
            old_file.unlink(missing_ok=True)
        with open(os.path.join(TWINS_DIR, f"{twin_id}.json"), "w") as f:
            json.dump(twin_data, f, indent=2)

    return {"twin_id": twin_id, "personality_model": personality_model}


def _compute_depth_score(data: dict) -> str:
    """Return 'Basic', 'Developed', or 'Deep' based on which of the 4 persona layers are populated.

    Layer 1 – Surface:        bio present (> 20 chars)
    Layer 2 – Decisions:      pastDecisions present
    Layer 3 – Values:         coreValues present
    Layer 4 – Meta-cognition: mindChange present (only filled by the deepen flow)
    """
    personality_model = data.get("personality_model") or {}
    ctx = personality_model.get("_context", {}) if isinstance(personality_model, dict) else {}
    core_values = personality_model.get("core_values") if isinstance(personality_model, dict) else None

    def _filled(key: str, min_len: int = 5) -> bool:
        v = ctx.get(key)
        return isinstance(v, str) and len(v.strip()) >= min_len

    def _values_layer_filled() -> bool:
        """Layer 3 – Values: consider either _context['coreValues'] or personality_model['core_values']."""
        # Prefer an explicit coreValues string in _context if present
        if _filled("coreValues"):
            return True
        # Fallback: check personality_model.core_values list
        if isinstance(core_values, list):
            for v in core_values:
                if isinstance(v, str) and v.strip():
                    return True
        return False

    layers = sum([
        _filled("bio", 20),
        _filled("pastDecisions"),
        _values_layer_filled(),
        _filled("mindChange"),
    ])

    # deepen_completed_at is stamped by _deepen_and_save whenever the interview
    # finishes. Use it as a reliable "Deep" signal because the LLM sometimes
    # covers a topic without populating the corresponding field_updates key,
    # leaving mindChange empty even though the user completed the flow.
    if data.get("deepen_completed_at") or layers == 4:
        return "Deep"
    if layers >= 2:
        return "Developed"
    return "Basic"


def _save_twin(twin_id: str, user_id: str, twin_data: dict) -> None:
    """Persist twin_data to both S3 keys (flat + per-user) or local disk."""
    if USE_S3:
        payload = json.dumps(twin_data, indent=2)
        for key in (
            f"{TWINS_S3_PREFIX}{twin_id}.json",
            f"{TWINS_S3_PREFIX}{user_id}/{twin_id}.json",
        ):
            s3_client.put_object(
                Bucket=S3_BUCKET,
                Key=key,
                Body=payload,
                ContentType="application/json",
            )
    else:
        os.makedirs(TWINS_DIR, exist_ok=True)
        with open(os.path.join(TWINS_DIR, f"{twin_id}.json"), "w") as f:
            json.dump(twin_data, f, indent=2)


def _save_twin_flat(twin_id: str, twin_data: dict) -> None:
    """Persist twin_data to only the flat twins/{twin_id}.json key.

    Used for background writes (e.g. the knowledge-gap ledger) that can be
    triggered by a chat turn against any twin, including ones with no owning
    user_id (public personas, the migrated default twin) — _save_twin's
    per-user key would break as twins/None/{twin_id}.json in that case. The
    per-user copy is read only by list_my_twins, a lightweight listing, so
    skipping it here just means that copy's knowledge_gaps can trail the
    flat copy slightly; nothing reads knowledge_gaps from the per-user copy.
    """
    if USE_S3:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=f"{TWINS_S3_PREFIX}{twin_id}.json",
            Body=json.dumps(twin_data, indent=2),
            ContentType="application/json",
        )
    else:
        os.makedirs(TWINS_DIR, exist_ok=True)
        with open(os.path.join(TWINS_DIR, f"{twin_id}.json"), "w") as f:
            json.dump(twin_data, f, indent=2)


def _delete_twin_storage(twin_id: str, user_id: str) -> None:
    """Delete both S3 keys (flat + per-user) or the local file for a twin."""
    if USE_S3:
        for key in (
            f"{TWINS_S3_PREFIX}{twin_id}.json",
            f"{TWINS_S3_PREFIX}{user_id}/{twin_id}.json",
        ):
            try:
                s3_client.delete_object(Bucket=S3_BUCKET, Key=key)
            except ClientError as e:
                print(f"Warning: failed to delete {key}: {e}")
    else:
        path = os.path.join(TWINS_DIR, f"{twin_id}.json")
        if os.path.exists(path):
            os.remove(path)


def _small_object_put(prefix: str, local_dir: str, subpath: str, data: dict) -> None:
    """Write a small JSON marker/record object — shares, public-share requests,
    featured markers. Mirrors _save_twin_flat's S3-vs-local branching."""
    payload = json.dumps(data, indent=2)
    if USE_S3:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=f"{prefix}{subpath}",
            Body=payload,
            ContentType="application/json",
        )
    else:
        full_path = os.path.join(local_dir, subpath)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, "w") as f:
            f.write(payload)


def _small_object_delete(prefix: str, local_dir: str, subpath: str) -> None:
    if USE_S3:
        try:
            s3_client.delete_object(Bucket=S3_BUCKET, Key=f"{prefix}{subpath}")
        except ClientError as e:
            print(f"Warning: failed to delete {prefix}{subpath}: {e}")
    else:
        full_path = os.path.join(local_dir, subpath)
        if os.path.exists(full_path):
            os.remove(full_path)


def _small_object_list(prefix: str, local_dir: str, sub_prefix: str = "") -> List[dict]:
    """List and parse all JSON objects under prefix + sub_prefix. Used for
    scoped listings (shares for one recipient email) and small unscoped
    listings (all pending public-share requests, all featured personas) —
    all expected to be low-volume, so a full listing under the prefix is fine."""
    results: List[dict] = []
    if USE_S3:
        paginator = s3_client.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=f"{prefix}{sub_prefix}"):
            for obj in page.get("Contents", []):
                try:
                    resp = s3_client.get_object(Bucket=S3_BUCKET, Key=obj["Key"])
                    results.append(json.loads(resp["Body"].read()))
                except Exception as e:
                    print(f"Warning: could not read {obj['Key']}: {e}")
    else:
        scan_dir = os.path.join(local_dir, sub_prefix)
        if os.path.isdir(scan_dir):
            for root, _dirs, files in os.walk(scan_dir):
                for fname in files:
                    if not fname.endswith(".json"):
                        continue
                    try:
                        with open(os.path.join(root, fname)) as f:
                            results.append(json.load(f))
                    except Exception as e:
                        print(f"Warning: could not read {fname}: {e}")
    return results


def _normalize_email(email: str) -> str:
    return email.strip().lower()


def _is_gap_worthy(orchestration: dict) -> bool:
    """Pure computation, no I/O — decides whether a chat turn's grounding
    signal is weak enough to be worth recording as a knowledge gap.

    True when the critic flagged the answer as uncertain or low-confidence,
    or when a factual/mixed query retrieved zero sources at all (advisory-only
    questions are expected to lean on inference per build_grounding_summary,
    so a thin advisory answer alone doesn't count)."""
    grounding = orchestration.get("grounding")
    query_type = orchestration.get("route", {}).get("query_type", "")
    retrieved_sources = orchestration.get("retrieved_sources") or []

    if grounding:
        if grounding.get("grounding_mode") == "uncertain":
            return True
        if grounding.get("confidence_label") == "low":
            return True
    if query_type in {"factual", "mixed"} and not retrieved_sources:
        return True
    return False


async def _record_chat_gap(twin_id: str, topic_tags: list, question_snippet: str) -> None:
    """Background task: load the twin fresh, record the gap, save it back.

    Reloading here (rather than reusing the twin_data already in the request's
    memory) narrows the race window on concurrent writes to the same twin —
    approximate counts under heavy concurrency are an accepted trade-off (see
    design.md), not a correctness guarantee this aims to provide.
    """
    try:
        twin_data = load_twin(twin_id)
        if not twin_data:
            return
        record_knowledge_gap(twin_data, topic_tags, source="inferred", question_snippet=question_snippet)
        _save_twin_flat(twin_id, twin_data)
    except Exception as exc:
        print(f"[gap-ledger] Failed to record gap for twin {twin_id}: {exc}")


class AddCorrectionRequest(BaseModel):
    question: str = Field(..., min_length=1, max_length=500)
    wrong_response: str = Field(..., min_length=1, max_length=500)
    correction: str = Field(..., min_length=1, max_length=500)

    @field_validator("question", "wrong_response", "correction", mode="before")
    @classmethod
    def strip_and_validate_non_empty(cls, v: Any) -> str:
        if v is None:
            raise ValueError("must not be empty")
        if not isinstance(v, str):
            raise TypeError("must be a string")
        v = v.strip()
        if not v:
            raise ValueError("must not be empty or whitespace")
        return v
@app.patch("/twin/{twin_id}/corrections")
async def add_correction(
    twin_id: str,
    request: AddCorrectionRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Append a user-supplied correction to a twin they own."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    corrections: list = twin_data.get("corrections", [])
    corrections.append({
        "question": request.question[:500],
        "wrong_response": request.wrong_response[:500],
        "correction": request.correction[:500],
        "created_at": datetime.now().isoformat(),
    })
    # Cap stored corrections to avoid unbounded growth
    twin_data["corrections"] = corrections[-20:]
    correction_source = build_correction_source(
        question=request.question[:500],
        wrong_response=request.wrong_response[:500],
        correction=request.correction[:500],
        created_at=corrections[-1]["created_at"],
    )
    twin_data["sources"] = merge_sources(
        twin_data.get("sources", []),
        [correction_source] if correction_source else [],
    )
    # Corrections are an explicit "the twin got this wrong" signal — stronger
    # than an inferred low-confidence chat turn, so they're recorded at
    # higher priority in the knowledge-gap ledger (see record_knowledge_gap's
    # source="correction" handling).
    record_knowledge_gap(
        twin_data,
        topic_tags_for_question(request.question),
        source="correction",
        question_snippet=request.question[:200],
    )
    _save_twin(twin_id, user_id, twin_data)
    return {"status": "ok", "corrections_count": len(twin_data["corrections"])}


# ---------------------------------------------------------------------------
# Persona vs Persona debate
# ---------------------------------------------------------------------------

# Number of rounds for the batched /chat/debate endpoint only.
# Each twin speaks once per round (total turns = DEBATE_ROUNDS * 2).
# Default kept at 2 to stay safely within the API Gateway 30s hard timeout —
# 4 sequential Bedrock calls at ~5-7s each ≈ 20-28s, leaving a safety margin.
# Set DEBATE_ROUNDS=3 via env var to enable 3 rounds (6 calls, ~24-30s — slim margin).
# Default kept at 3 to match the frontend NEXT_PUBLIC_DEBATE_ROUNDS default
# (3 rounds, 6 total Bedrock calls). This can bring total latency close to
# the API Gateway 30s hard timeout (6 calls at ~5-7s each ≈ 30-42s).
# If you observe timeouts, lower DEBATE_ROUNDS to 2 via env var to stay more
# safely within the timeout (4 calls, ~20-28s total).
# NOTE: This does NOT govern /debate/turn (turn-by-turn). That endpoint handles
# a single turn per request; the number of turns is driven entirely by the
# frontend's NEXT_PUBLIC_DEBATE_ROUNDS env var (default 3). Both env vars
# should be set to the same value to keep the two debate modes consistent.
_DEBATE_ROUNDS_DEFAULT = 3
_DEBATE_ROUNDS_MIN = 1
_DEBATE_ROUNDS_MAX = 3
_debate_rounds_raw = os.getenv("DEBATE_ROUNDS", "").strip()
try:
    _debate_rounds_val = int(_debate_rounds_raw) if _debate_rounds_raw else _DEBATE_ROUNDS_DEFAULT
except (TypeError, ValueError):
    _debate_rounds_val = _DEBATE_ROUNDS_DEFAULT
if _debate_rounds_val < _DEBATE_ROUNDS_MIN:
    _debate_rounds_val = _DEBATE_ROUNDS_MIN
elif _debate_rounds_val > _DEBATE_ROUNDS_MAX:
    _debate_rounds_val = _DEBATE_ROUNDS_MAX
DEBATE_ROUNDS = _debate_rounds_val


class DebateAgent:
    """Autonomous agent representing a single twin in a structured debate.

    Each agent maintains its own conversation history so it has independent,
    persona-consistent context across the exchange without seeing the other
    twin's internal state.
    """

    def __init__(self, twin_data: dict) -> None:
        self.twin_id: str = twin_data["twin_id"]
        self.name: str = twin_data.get("name", "Unknown")
        self.title: str = twin_data.get("title", "")
        personality_model = twin_data.get("personality_model", {})
        self._system_prompt: str = prompt(
            personality_model=personality_model,
            twin_name=self.name,
            twin_title=self.title,
        )
        self._history: List[Dict] = []  # agent's own view of the debate

    def _build_messages(self, user_turn: str) -> List[Dict]:
        messages: List[Dict] = [
            {"role": "user", "content": [{"text": f"System: {self._system_prompt}"}]}
        ]
        for msg in self._history[-10:]:  # cap at last 10 messages (5 exchanges)
            messages.append({"role": msg["role"], "content": [{"text": msg["content"]}]})
        messages.append({"role": "user", "content": [{"text": user_turn}]})
        return messages

    def respond(self, user_turn: str) -> str:
        """Call Bedrock synchronously, update internal history, return response text.

        Designed to be called via asyncio.to_thread so it doesn't block the
        event loop during the debate orchestration.
        """
        messages = self._build_messages(user_turn)
        response = bedrock_client.converse(
            modelId=BEDROCK_MODEL_ID,
            messages=messages,
            inferenceConfig={"maxTokens": 200, "temperature": 0.75, "topP": 0.9},
        )
        text: str = response["output"]["message"]["content"][0]["text"]
        self._history.append({"role": "user", "content": user_turn})
        self._history.append({"role": "assistant", "content": text})
        return text


# Maximum allowed number of history entries for /debate/turn.
# Intentionally decoupled from DEBATE_ROUNDS / frontend NEXT_PUBLIC_DEBATE_ROUNDS
# so config drift cannot cause mid-debate 422s. Can be overridden via env var.
try:
    _MAX_HISTORY_ENTRIES = int(os.getenv("DEBATE_MAX_HISTORY_ENTRIES", "20"))
except (TypeError, ValueError):
    _MAX_HISTORY_ENTRIES = 20
if _MAX_HISTORY_ENTRIES < 1:
    _MAX_HISTORY_ENTRIES = 1
_MAX_TWIN_NAME_LEN = 100
# 1000 chars per entry: generous for 3-5 sentences (~300-500 chars typical).
_MAX_HISTORY_TEXT_LEN = 1000
# Total character budget for history injected into the prompt.
# Oldest entries are dropped server-side if the budget is exceeded.
_MAX_HISTORY_TOTAL_CHARS = 8000


class DebateHistoryEntry(BaseModel):
    twin_name: str
    text: str

    @field_validator("twin_name")
    @classmethod
    def twin_name_length(cls, v: str) -> str:
        if len(v) > _MAX_TWIN_NAME_LEN:
            raise ValueError(f"twin_name must be {_MAX_TWIN_NAME_LEN} characters or fewer")
        return v

    @field_validator("text")
    @classmethod
    def text_length(cls, v: str) -> str:
        if len(v) > _MAX_HISTORY_TEXT_LEN:
            raise ValueError(f"history text must be {_MAX_HISTORY_TEXT_LEN} characters or fewer")
        return v


class DebateTurnRequest(BaseModel):
    twin_id: str
    topic: str
    history: List[DebateHistoryEntry] = Field(default_factory=list)  # full debate so far, oldest first

    @field_validator("history")
    @classmethod
    def history_max_entries(cls, v: List[DebateHistoryEntry]) -> List[DebateHistoryEntry]:
        if len(v) > _MAX_HISTORY_ENTRIES:
            raise ValueError(f"history must not exceed {_MAX_HISTORY_ENTRIES} entries")
        return v

    @field_validator("topic")
    @classmethod
    def topic_not_empty(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("topic must not be empty")
        if len(v) > 500:
            raise ValueError("topic must be 500 characters or fewer")
        return v


class DebateRequest(BaseModel):
    twin_id_a: str
    twin_id_b: str
    topic: str

    @field_validator("topic")
    @classmethod
    def topic_not_empty(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("topic must not be empty")
        if len(v) > 500:
            raise ValueError("topic must be 500 characters or fewer")
        return v


class DebateTurn(BaseModel):
    twin_id: str
    twin_name: str
    turn_number: int
    text: str


class DebateResponse(BaseModel):
    topic: str
    turns: List[DebateTurn]


class DebateTurnResponse(BaseModel):
    twin_id: str
    twin_name: str
    text: str


def _is_debate_authorized(twin_data: dict, user_id: str) -> bool:
    """A twin may be used in a debate if the caller owns it, or if it has no
    owner at all — public personas and the default twin (before
    DEFAULT_TWIN_OWNER_USER_ID is configured) never have a user_id, and are
    meant to be usable by any signed-in user as a debate partner."""
    owner = twin_data.get("user_id")
    return owner is None or owner == user_id


@app.post("/debate/turn", response_model=DebateTurnResponse)
async def debate_turn(
    request: DebateTurnRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Generate a single debate turn for one twin.

    The frontend drives the debate loop: it calls this endpoint once per turn,
    passing the full history so far. This makes each agent's response feel live
    (typing indicator while waiting, typewriter animation on arrival) without
    requiring Lambda response streaming.
    """
    twin_data = load_twin(request.twin_id)
    if not twin_data or not _is_debate_authorized(twin_data, user_id):
        raise HTTPException(status_code=404, detail="Twin not found")

    agent = DebateAgent(twin_data)

    # Build the debate-context prompt from history
    def _esc(s: str) -> str:
        """JSON-escape a string so quotes/newlines don't break prompt structure."""
        return json.dumps(s)[1:-1]

    if not request.history:
        turn_prompt = (
            f'You are in a live debate on the topic: "{_esc(request.topic)}". '
            f"Open with your perspective. Speak in your natural voice. 3-5 sentences."
        )
    else:
        # Server-side truncation: drop oldest entries until total chars fit in budget.
        history = list(request.history)
        total_chars = sum(len(e.twin_name) + len(e.text) for e in history)
        while len(history) > 1 and total_chars > _MAX_HISTORY_TOTAL_CHARS:
            dropped = history.pop(0)
            total_chars -= len(dropped.twin_name) + len(dropped.text)

        history_lines = "\n".join(
            f'{_esc(e.twin_name)}: "{_esc(e.text)}"' for e in history
        )
        last = history[-1]
        turn_prompt = (
            f'You are in a live debate on the topic: "{_esc(request.topic)}".\n\n'
            f"Debate so far:\n{history_lines}\n\n"
            f'{_esc(last.twin_name)} just said: "{_esc(last.text)}"\n\n'
            f"Respond to their point. Stay in character. 3-5 sentences."
        )

    try:
        text = await asyncio.to_thread(agent.respond, turn_prompt)
    except ClientError as e:
        print(f"Bedrock error in debate/turn: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate response")
    except Exception as e:
        print(f"Unexpected error in debate/turn: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate response")

    return {"twin_id": agent.twin_id, "twin_name": agent.name, "text": text}


@app.post("/chat/debate", response_model=DebateResponse)
async def debate(
    request: DebateRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Run a structured debate between two twins — either the caller's own,
    or a public persona (which has no owner and is usable by any signed-in user).

    Each twin is instantiated as an independent DebateAgent with its own
    persona and conversation context. The orchestrator alternates calls for
    DEBATE_ROUNDS rounds (each twin speaks DEBATE_ROUNDS times = 2×DEBATE_ROUNDS
    total turns).

    Note: responses are buffered and returned as a single JSON payload.
    True token-level streaming requires Lambda Function URL response streaming
    and is a planned future upgrade.
    """
    twin_a_data = load_twin(request.twin_id_a)
    twin_b_data = load_twin(request.twin_id_b)

    if not twin_a_data or not _is_debate_authorized(twin_a_data, user_id):
        raise HTTPException(status_code=404, detail="Twin A not found")
    if not twin_b_data or not _is_debate_authorized(twin_b_data, user_id):
        raise HTTPException(status_code=404, detail="Twin B not found")
    if request.twin_id_a == request.twin_id_b:
        raise HTTPException(status_code=400, detail="Debate requires two different twins")
    # Public personas have no owner and are usable by anyone signed in, but at
    # least one side must be the caller's own twin — otherwise any two signed-in
    # users could rack up public-persona-vs-public-persona debates with nothing
    # of their own involved, which isn't the intended "debate your own twin
    # against a public persona" use case.
    if twin_a_data.get("user_id") != user_id and twin_b_data.get("user_id") != user_id:
        raise HTTPException(status_code=400, detail="At least one twin in the debate must be your own")

    agent_a = DebateAgent(twin_a_data)
    agent_b = DebateAgent(twin_b_data)

    turns: List[Dict] = []
    last_text = ""

    try:
        for round_num in range(DEBATE_ROUNDS):
            # ── Agent A's turn ────────────────────────────────────────────
            if round_num == 0:
                prompt_a = (
                    f'You are debating {agent_b.name} on the topic: "{request.topic}". '
                    f"Open with your perspective. Be direct and speak in your natural voice. "
                    f"Keep it to 3-5 sentences."
                )
            else:
                prompt_a = (
                    f'{agent_b.name} said: "{last_text}"\n\n'
                    f"Respond to their point in the debate. Stay in character. "
                    f"Keep it to 3-5 sentences."
                )
            response_a = await asyncio.to_thread(agent_a.respond, prompt_a)
            turns.append({"twin_id": agent_a.twin_id, "twin_name": agent_a.name,
                          "turn_number": len(turns) + 1, "text": response_a})
            last_text = response_a

            # ── Agent B's turn ────────────────────────────────────────────
            if round_num == 0:
                prompt_b = (
                    f'You are debating {agent_a.name} on the topic: "{request.topic}". '
                    f'{agent_a.name} just said: "{response_a}"\n\n'
                    f"Respond with your perspective. Be direct and speak in your natural voice. "
                    f"Keep it to 3-5 sentences."
                )
            else:
                prompt_b = (
                    f'{agent_a.name} said: "{last_text}"\n\n'
                    f"Respond to their point in the debate. Stay in character. "
                    f"Keep it to 3-5 sentences."
                )
            response_b = await asyncio.to_thread(agent_b.respond, prompt_b)
            turns.append({"twin_id": agent_b.twin_id, "twin_name": agent_b.name,
                          "turn_number": len(turns) + 1, "text": response_b})
            last_text = response_b

    except ClientError as e:
        print(f"Bedrock ClientError in debate: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate debate response")
    except Exception as e:
        print(f"Unexpected error in debate: {e}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred during the debate")

    return {"topic": request.topic, "turns": turns}


# ---------------------------------------------------------------------------
# Chat-based onboarding interview
# ---------------------------------------------------------------------------

_ONBOARD_SYSTEM_TEMPLATE = """\
You are a sharp, warm interviewer helping someone build their AI twin — a digital version of them \
that can answer questions on their behalf.

Your job: learn who they are through a natural conversation, covering 6 topics in order.

TOPICS:
1. IDENTITY      → name, job title, short bio (who they are in a sentence or two)
2. PROFESSIONAL  → key skills, career story, one notable achievement{linkedin_skip}
3. DECISIONS     → how they make hard calls (want a real example), risk appetite
4. VALUES        → what they stand for; one thing they would push back on under pressure
5. WORKING_STYLE → how they communicate; what colleagues sometimes misread about them
6. VOICE         → a phrase they overuse or a verbal tic; bullet-points or paragraphs?

RULES — follow exactly:
- One question per turn. 2–4 sentences per message. Be concise.
- CRITICAL: Every message MUST end with a question, except the final closing message \
when done is true. Never just acknowledge — always ask about the next remaining topic \
in the same message. If you acknowledge, do it in one short phrase, then immediately ask.
- If an answer is vague or generic (e.g. "I just go with my gut", "I value honesty"), \
push back ONCE: invent a tiny relatable story in one sentence that mirrors the vague answer \
(first-person or "I once worked with someone who..."), then re-ask more concretely. \
Push back only once per topic — then accept and move on regardless of the answer.
- After a rich or interesting answer, acknowledge in one brief phrase \
("Got it.", "That's clear.", "Interesting.") and immediately ask the next question.
- Mirror their tone: terse answers → short questions; expressive answers → slightly warmer.
- Never use form-speak ("Question 3 of 6", "Next section", "Moving on to topic...").
- When all 6 topics are covered (none remaining), close with one natural sentence and set done to true.

CURRENT STATE:
Topics remaining: {topics_remaining}
Fields collected so far:
{fields_json}
{linkedin_section}

RETURN ONLY valid JSON — no markdown, no text outside the JSON object.
NEVER include JSON or curly-brace fragments inside the "message" string value.
The "done" field belongs only in the top-level JSON structure, not in the message text:
{{
  "message": "your conversational response and next question as natural prose",
  "field_updates": {{
    "name": "value or omit key entirely if not in this message",
    "title": "...",
    "bio": "...",
    "skills": "...",
    "experience": "...",
    "achievements": "...",
    "coreValues": "...",
    "decisionStyle": "...",
    "riskTolerance": "low or medium or high — omit if unclear",
    "pastDecisions": "...",
    "communicationStyle": "...",
    "blindSpots": "...",
    "verbalQuirks": "...",
    "responseStyle": "concise or balanced or detailed — omit if unclear"
  }},
  "topics_covered": ["IDENTITY", "PROFESSIONAL"],
  "done": false
}}

When done is true, set "done": true in the JSON above. Do NOT include a twin_payload field — \
the client will assemble the twin from the collected fields_collected data.
"""


class OnboardHistoryItem(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class OnboardRequest(BaseModel):
    history: List[OnboardHistoryItem] = Field(default_factory=list)
    linkedin_parsed: Optional[Dict[str, Any]] = None
    fields_collected: Optional[Dict[str, Any]] = None
    topics_covered: List[str] = Field(default_factory=list)


class OnboardFieldUpdates(BaseModel):
    """Validated field updates extracted from the model response."""

    name: Optional[str] = None
    title: Optional[str] = None
    bio: Optional[str] = None
    skills: Optional[str] = None
    experience: Optional[str] = None
    achievements: Optional[str] = None
    coreValues: Optional[str] = None
    decisionStyle: Optional[str] = None
    riskTolerance: Optional[str] = None
    pastDecisions: Optional[str] = None
    communicationStyle: Optional[str] = None
    blindSpots: Optional[str] = None
    verbalQuirks: Optional[str] = None
    responseStyle: Optional[str] = None

    model_config = {"extra": "ignore"}


class OnboardResponse(BaseModel):
    """Validated response returned by /onboard/message."""

    message: str
    field_updates: OnboardFieldUpdates = Field(default_factory=OnboardFieldUpdates)
    topics_covered: List[str] = Field(default_factory=list)
    done: bool = False
    twin_payload: Optional[Dict[str, Any]] = None

    model_config = {"extra": "ignore"}

    @field_validator("field_updates", mode="before")
    @classmethod
    def _coerce_field_updates(cls, v: Any) -> Any:
        """Accept a dict or fall back to an empty OnboardFieldUpdates."""
        if not isinstance(v, dict):
            return {}
        return v

    @field_validator("topics_covered", mode="before")
    @classmethod
    def _coerce_topics_covered(cls, v: Any) -> Any:
        """Accept a list of strings or fall back to an empty list."""
        if not isinstance(v, list):
            return []
        return [item for item in v if isinstance(item, str)]

    @field_validator("done", mode="before")
    @classmethod
    def _coerce_done(cls, v: Any) -> Any:
        """Accept a bool; coerce any non-bool value to False."""
        if isinstance(v, bool):
            return v
        return False


_ALL_ONBOARD_TOPICS = ["IDENTITY", "PROFESSIONAL", "DECISIONS", "VALUES", "WORKING_STYLE", "VOICE"]


@app.post("/onboard/message")
async def onboard_message(
    request: OnboardRequest,
    _user_id: str = Depends(get_current_user_id),
):
    covered = list(request.topics_covered)
    # Auto-mark PROFESSIONAL covered when LinkedIn data is provided
    if request.linkedin_parsed and "PROFESSIONAL" not in covered:
        covered.append("PROFESSIONAL")
    # Auto-mark IDENTITY covered when name+title+bio are already collected
    # (from LinkedIn, a previous turn, or any other source)
    fields = request.fields_collected or {}
    if (
        "IDENTITY" not in covered
        and fields.get("name")
        and fields.get("title")
        and fields.get("bio")
    ):
        covered.append("IDENTITY")

    remaining = [t for t in _ALL_ONBOARD_TOPICS if t not in covered]

    linkedin_skip = " (SKIP — LinkedIn PDF provided)" if request.linkedin_parsed else ""

    linkedin_section = ""
    if request.linkedin_parsed:
        lp = request.linkedin_parsed
        lines = []
        if lp.get("name"):        lines.append(f"Name: {lp['name']}")
        if lp.get("title"):       lines.append(f"Title: {lp['title']}")
        if lp.get("skills"):      lines.append(f"Skills: {str(lp['skills'])[:300]}")
        if lp.get("experience"):  lines.append(f"Experience: {str(lp['experience'])[:400]}")
        if lp.get("achievements"): lines.append(f"Achievements: {str(lp['achievements'])[:200]}")
        if lines:
            linkedin_section = "\nLinkedIn PDF already parsed (use this, don't re-ask):\n" + "\n".join(lines)

    system_prompt = _ONBOARD_SYSTEM_TEMPLATE.format(
        linkedin_skip=linkedin_skip,
        topics_remaining=", ".join(remaining) if remaining else "ALL COVERED",
        fields_json=json.dumps(request.fields_collected or {}, indent=2),
        linkedin_section=linkedin_section,
    )

    messages: List[Dict[str, Any]] = []
    if not request.history:
        # Seed with a minimal opener so the model produces the first question
        messages = [{"role": "user", "content": [{"text": "hi, let's start"}]}]
    else:
        # Cap the amount of history sent to Bedrock to avoid unbounded prompts
        for item in request.history[-50:]:
            if item.role == "user":
                role = "user"
            elif item.role == "assistant":
                role = "assistant"
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid role in history item: {item.role!r}. Allowed roles are 'user' and 'assistant'.",
                )
            messages.append({"role": role, "content": [{"text": item.content}]})
        # Bedrock requires the conversation to start with a user turn.
        # When history only contains the bot's opening question (assistant), prepend
        # the synthetic opener so the first message is always from "user".
        if messages and messages[0]["role"] != "user":
            messages.insert(0, {"role": "user", "content": [{"text": "hi, let's start"}]})

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            system=[{"text": system_prompt}],
            messages=messages,
            inferenceConfig={"maxTokens": 700, "temperature": 0.9, "topP": 0.95},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()

        # Use robust JSON extraction — handles code fences and leading/trailing text.
        # Pass required_key so stray {"done": true} fragments are skipped.
        data = _extract_json_object(raw, required_key="message")

        if not isinstance(data, dict) or "message" not in data:
            raise ValueError("Invalid onboarding JSON structure")

        # Validate and coerce model output against a strict response model so that
        # missing keys default safely and wrong types don't reach the frontend.
        validated = OnboardResponse.model_validate(data)
        return validated.model_dump(exclude_none=True)
    except (ValueError, json.JSONDecodeError) as exc:
        print(f"Onboard JSON parse error: {exc!r}")
        if os.getenv("DEBUG_LOG_ONBOARD_RAW") == "1":
            print(f"Onboard raw snippet (truncated): {raw[:200]!r}")

        # Salvage a clean plain-text message and detect the done signal.
        # Strategy: look for a trailing JSON fragment (rfind '{' in the latter half
        # of the text) and parse it with json.loads — only trust done:true when it
        # appears as an actual top-level key in a parseable object, not anywhere in
        # the raw string (which would produce false positives on quoted examples).
        done_msg = "Thanks — that's everything I need! Let me put your twin together."
        cont_msg = "Got it — let me keep going. Could you tell me a bit more?"

        fallback_done = False
        fallback_message = raw

        if raw.strip().startswith("{") or raw.strip().startswith("```"):
            # Entire output looks like a (possibly malformed) JSON blob — can't
            # salvage natural language, so use a canned message.
            fallback_message = cont_msg
        else:
            # Check for a trailing JSON fragment like  ...nice. {"done": true}
            last_brace = raw.rfind('{')
            if last_brace > len(raw) // 2:
                try:
                    fragment = json.loads(raw[last_brace:])
                    # Only treat this as a real trailer (and truncate) if it parses
                    # and looks like the expected {"done": ...} object.
                    if isinstance(fragment, dict) and "done" in fragment:
                        if fragment.get("done") is True:
                            fallback_done = True
                        fallback_message = raw[:last_brace].strip()
                except json.JSONDecodeError:
                    # Leave fallback_message as the full raw text if the fragment
                    # doesn't parse; don't truncate on malformed JSON.
                    pass
            if not fallback_message:
                fallback_message = done_msg if fallback_done else cont_msg

        fallback_topics = list(_ALL_ONBOARD_TOPICS) if fallback_done else covered

        return {
            "message": fallback_message,
            "field_updates": {},
            "topics_covered": fallback_topics,
            "done": fallback_done,
        }
    except Exception as exc:
        print(f"Unexpected error in /onboard/message: {exc}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred. Please try again.")


_DEEPEN_TOPIC_SYSTEM_TEMPLATE = """\
You are interviewing someone to deepen one specific trait of their AI twin's personality model — \
uncovering the nuance that makes reasoning feel real instead of generic.

CURRENT TOPIC: {topic_hint}

Existing twin context (for reference — do NOT repeat back to them):
{existing_context}

Your job this turn:
- If there's no prior conversation on this topic yet, ask an engaging opening question based on the \
topic above — adapt it into your own natural phrasing, don't read it verbatim.
- Otherwise, ask ONE natural, specific follow-up that would surface something concrete you don't have \
yet. Don't repeat a question you already asked. Push for specifics over platitudes.
- Do NOT judge or acknowledge whether the answer is "enough" — a separate process decides that. Just \
keep the conversation moving forward with a good next question, as if you don't yet know whether \
you'll need to ask more.
- Extract the user's actual answer content from their latest message into "answer_delta" — their \
words/content, not your own questions. Leave it empty if there's no new user answer yet (e.g. this \
is the opening question).
- 2-4 sentences max. Be direct and warm. Never sound like a form ("Question 1 of 3").

RETURN ONLY valid JSON — no markdown, no text outside the JSON:
{{
  "message": "your question, as natural prose",
  "answer_delta": "the user's new answer content this turn, or empty string"
}}
"""

_DEEPEN_CRITIC_MODEL_ID = "amazon.nova-micro-v1:0"
_DEEPEN_CRITIC_PROMPT = (
    "You are judging whether an interview answer is specific and concrete enough to meaningfully "
    "update someone's personality profile, as opposed to vague, generic, or incomplete.\n\n"
    "Topic being explored: {topic_hint}\n\n"
    "Accumulated answer so far:\n{answer}\n\n"
    "Reply in exactly this format: YES|<one-sentence reason> or NO|<what's still missing>"
)


async def _judge_topic_sufficiency(topic_hint: str, answer: str) -> tuple[bool, str]:
    """Cheap-model critic call deciding whether the accumulated answer for the
    current deepen topic is concrete enough to save, or needs another
    follow-up. Short-circuits on trivially-empty answers to avoid a wasted
    Bedrock call on the topic's opening turn."""
    if not answer or len(answer.strip()) < 15:
        return False, "No substantive answer yet."
    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=_DEEPEN_CRITIC_MODEL_ID,
            messages=[{"role": "user", "content": [
                {"text": _DEEPEN_CRITIC_PROMPT.format(topic_hint=topic_hint, answer=answer)}
            ]}],
            inferenceConfig={"maxTokens": 60, "temperature": 0},
        )
        text = response["output"]["message"]["content"][0]["text"].strip()
        is_sufficient = text.upper().startswith("YES")
        reason = text.split("|", 1)[1].strip() if "|" in text else ""
        return is_sufficient, reason
    except Exception as exc:
        print(f"[deepen] Critic judgment failed, defaulting to insufficient: {exc}")
        return False, ""

# ── Deepen topic library ──────────────────────────────────────────────────
# Static topics (evergreen fallback + thin-field prompts) with known question
# hints and which personality_model fields they should sharpen. GAP:* topics
# are dynamic (derived from a twin's knowledge_gaps ledger at request time)
# and aren't in this dict — see _target_fields_for_topics' GAP: branch.
_EVERGREEN_DEEPEN_TOPICS: Dict[str, Dict[str, Any]] = {
    "PAST_DECISIONS": {
        "question_hint": (
            "Walk me through 2-3 decisions you've made that were genuinely hard — "
            "what you chose, what you gave up, and whether you'd do it again."
        ),
        "target_fields": ["decision_heuristics", "blind_spots", "decision_framework", "pivotal_decisions"],
        "field_update_key": "pastDecisions",
    },
    "NON_NEGOTIABLES": {
        "question_hint": (
            "What would you flat-out refuse to do even under real pressure? "
            "And what would you bend on if the trade-off was right?"
        ),
        "target_fields": ["blind_spots", "what_they_avoid", "decision_framework"],
        "field_update_key": "nonNegotiables",
    },
    "MIND_CHANGE": {
        "question_hint": (
            "Tell me about a time you changed your mind on something you'd held "
            "strongly. What actually moved you?"
        ),
        "target_fields": ["mind_change", "personality_summary"],
        "field_update_key": "mindChange",
    },
}

_THIN_FIELD_DEEPEN_TOPICS: Dict[str, Dict[str, Any]] = {
    "FIELD:pivotal_decisions": {
        "question_hint": "Tell me about 1-2 pivotal moments or turning-point decisions in your career or life.",
        "target_fields": ["pivotal_decisions"],
        "min_items": 2,
    },
    "FIELD:characteristic_quotes": {
        "question_hint": (
            "Give me a couple of things you actually say often, in your own words — "
            "phrases or lines that sound like you."
        ),
        "target_fields": ["characteristic_quotes"],
        "min_items": 2,
    },
    "FIELD:core_values": {
        "question_hint": (
            "What 3-5 values genuinely drive your decisions — not aspirational ones, "
            "ones you actually act on when it's inconvenient to?"
        ),
        "target_fields": ["core_values"],
        "min_items": 3,
    },
    "FIELD:what_they_optimize_for": {
        "question_hint": "When you're weighing a real tradeoff, what are you actually optimizing for?",
        "target_fields": ["what_they_optimize_for"],
        "min_items": 3,
    },
    "FIELD:what_they_avoid": {
        "question_hint": "What do you actively steer away from in how you work or decide things?",
        "target_fields": ["what_they_avoid"],
        "min_items": 3,
    },
    "FIELD:communication_traits": {
        "question_hint": "How would people who've worked closely with you describe the way you communicate?",
        "target_fields": ["communication_traits"],
        "min_items": 3,
    },
    "FIELD:blind_spots": {
        "question_hint": "What's a blind spot or weakness you're aware of in how you work or decide things?",
        "target_fields": ["blind_spots"],
        "min_items": 2,
    },
    "FIELD:risk_profile": {
        "question_hint": "How do you actually behave under risk or uncertainty — cautious, calculated, or comfortable with ambiguity?",
        "target_fields": ["risk_profile"],
        "min_items": 1,
    },
}

_GAP_TOPIC_DEFAULT_TARGET_FIELDS = ["decision_heuristics", "blind_spots", "personality_summary"]
_MAX_DEEPEN_TOPICS_PER_SESSION = 3
_CONSOLIDATION_EVERY_N_PATCHES = 5
_MODEL_VERSIONS_CAP = 10


def _is_field_thin(personality_model: dict, field: str, min_items: int = 2) -> bool:
    val = personality_model.get(field)
    if not val:
        return True
    if isinstance(val, list):
        return len(val) < min_items
    if isinstance(val, str):
        return len(val.strip()) < 20
    return False


def _select_deepen_topics(
    twin_data: dict, covered: List[str], limit: int = _MAX_DEEPEN_TOPICS_PER_SESSION
) -> List[Dict[str, Any]]:
    """Priority chain: gap ledger -> research-agent candidates -> thin
    personality_model fields -> evergreen fallback bank. Returns up to
    `limit` topics not already in `covered` and not otherwise already
    answered (evergreen/thin-field topics re-check their underlying field's
    actual state, not just `covered`, so this is self-correcting across
    sessions rather than relying on the caller to track it). A twin with
    nothing left across all four tiers returns [] — this never writes a
    permanent "fully deepened" marker; the next call simply recomputes
    against whatever's changed since (new gaps, edited fields, a fresh
    research-agent pass, etc.)."""
    pm = twin_data.get("personality_model", {}) or {}
    ctx = pm.get("_context", {}) if isinstance(pm, dict) else {}
    covered_set = set(covered)
    selected: List[Dict[str, Any]] = []

    gaps = sorted(
        (twin_data.get("knowledge_gaps") or []),
        key=lambda g: int(g.get("count", 0)),
        reverse=True,
    )
    for gap in gaps:
        if len(selected) >= limit:
            break
        if int(gap.get("count", 0)) <= 0:
            continue
        tags = gap.get("topic_tags") or []
        topic_id = "GAP:" + "|".join(sorted(tags))
        if topic_id in covered_set:
            continue
        tag_text = ", ".join(tags) if tags else "a topic"
        selected.append({
            "id": topic_id,
            "question_hint": (
                f"People keep asking about {tag_text} and the twin's answers have been weak or "
                f"ungrounded — ask a focused question that would give the twin real material here."
            ),
            "target_fields": _GAP_TOPIC_DEFAULT_TARGET_FIELDS,
        })

    if len(selected) < limit:
        # Novel angles proposed by the periodic research agent (see
        # run_research_agent) — curated, persona-specific questions beyond
        # the fixed taxonomy below. Consumed (removed) once answered and
        # saved, via _consume_research_candidates.
        for candidate in (twin_data.get("research_candidate_topics") or []):
            if len(selected) >= limit:
                break
            topic_id = candidate.get("id", "")
            if not topic_id or topic_id in covered_set:
                continue
            selected.append({
                "id": topic_id,
                "question_hint": candidate.get("question_hint", ""),
                "target_fields": candidate.get("target_fields") or _GAP_TOPIC_DEFAULT_TARGET_FIELDS,
            })

    if len(selected) < limit:
        for topic_id, meta in _THIN_FIELD_DEEPEN_TOPICS.items():
            if len(selected) >= limit:
                break
            if topic_id in covered_set:
                continue
            if any(_is_field_thin(pm, f, meta.get("min_items", 2)) for f in meta["target_fields"]):
                selected.append({
                    "id": topic_id,
                    "question_hint": meta["question_hint"],
                    "target_fields": meta["target_fields"],
                })

    if len(selected) < limit:
        for topic_id, meta in _EVERGREEN_DEEPEN_TOPICS.items():
            if len(selected) >= limit:
                break
            if topic_id in covered_set:
                continue
            field_key = meta.get("field_update_key")
            if field_key and isinstance(ctx, dict) and ctx.get(field_key):
                continue  # already answered in a prior session
            selected.append({
                "id": topic_id,
                "question_hint": meta["question_hint"],
                "target_fields": meta["target_fields"],
            })

    return selected


def _target_fields_for_topics(topic_ids: List[str], twin_data: Optional[dict] = None) -> List[str]:
    fields: set[str] = set()
    research_candidates = {
        c["id"]: c for c in (twin_data.get("research_candidate_topics") or [])
    } if twin_data else {}
    for tid in topic_ids:
        if tid in _EVERGREEN_DEEPEN_TOPICS:
            fields.update(_EVERGREEN_DEEPEN_TOPICS[tid]["target_fields"])
        elif tid in _THIN_FIELD_DEEPEN_TOPICS:
            fields.update(_THIN_FIELD_DEEPEN_TOPICS[tid]["target_fields"])
        elif tid.startswith("GAP:"):
            fields.update(_GAP_TOPIC_DEFAULT_TARGET_FIELDS)
        elif tid.startswith("RESEARCH:") and tid in research_candidates:
            fields.update(research_candidates[tid].get("target_fields") or _GAP_TOPIC_DEFAULT_TARGET_FIELDS)
    return sorted(fields) if fields else list(_GAP_TOPIC_DEFAULT_TARGET_FIELDS)


def _decay_addressed_gaps(twin_data: dict, addressed_topic_ids: List[str]) -> None:
    """Reset the count of any gap-ledger entry this session directly
    addressed, so it stops dominating topic selection until it re-accumulates
    from new chat activity (task 1.5)."""
    addressed_tag_sets = {
        tuple(sorted(tid[len("GAP:"):].split("|")))
        for tid in addressed_topic_ids
        if tid.startswith("GAP:")
    }
    if not addressed_tag_sets:
        return
    for gap in (twin_data.get("knowledge_gaps") or []):
        tags_key = tuple(sorted(gap.get("topic_tags", [])))
        if tags_key in addressed_tag_sets:
            gap["count"] = 0


def _consume_research_candidates(twin_data: dict, addressed_topic_ids: List[str]) -> None:
    """Remove a research-agent candidate topic once it's been answered and
    saved, so it doesn't linger in the backlog alongside the now-real data."""
    addressed = {tid for tid in addressed_topic_ids if tid.startswith("RESEARCH:")}
    if not addressed:
        return
    candidates = twin_data.get("research_candidate_topics") or []
    twin_data["research_candidate_topics"] = [c for c in candidates if c.get("id") not in addressed]


# ── Research agent (periodic, not per-request) ──────────────────────────────
# Analyzes a twin's full profile and proposes novel deepen-interview angles
# beyond the fixed taxonomy above — the "what haven't we thought to ask"
# half of the deepen redesign, run on a schedule (see run_research_agent_batch
# and the twice-monthly EventBridge Scheduler in terraform) rather than on
# every chat/deepen request, since it's exploratory and doesn't need to be
# fresh turn-to-turn.
_MAX_RESEARCH_CANDIDATE_TOPICS = 15

_RESEARCH_AGENT_PROMPT_TEMPLATE = """You are a research strategist helping build the most complete, \
accurate AI twin possible for {name}{title_suffix}.

You MUST propose exactly 3 to 5 new interview angles, even if the current profile looks reasonably \
complete — there is always something more specific or surprising worth exploring for a real person. \
Do not return an empty list under any circumstance.

Study everything captured about them so far:
{personality_model_json}

RECENT KNOWLEDGE GAPS (topics people have asked the twin about where its answers were weak):
{knowledge_gaps_summary}

ALREADY-PROPOSED CANDIDATE TOPICS (do not repeat these or anything substantially similar):
{existing_candidates_summary}

Propose angles that go beyond generic questions — tie each one to specifics already in the profile \
above, not a template. Example of a GOOD angle for someone whose decision framework mentions weighing \
team impact heavily: "Describe a time you had to choose between a teammate's growth opportunity and \
hitting a deadline — what did you do?" Example of a BAD angle (too generic, don't do this): "What are \
your values?"

Return ONLY a valid JSON array, no markdown, no other text:
[
  {{"question_hint": "a natural, specific interview question", "target_fields": ["personality_model field(s) this would sharpen"], "rationale": "one sentence on why this angle matters for this specific person"}}
]
"""


async def run_research_agent(twin_id: str) -> int:
    """Analyze one twin and append new candidate deepen topics. Returns the
    number of candidates added (0 if the twin already has a healthy backlog,
    the model call failed, or nothing usable was returned — all non-fatal,
    since this always runs as part of a best-effort batch, never inline in a
    user-facing request)."""
    twin_data = load_twin(twin_id)
    if not twin_data:
        return 0
    pm = twin_data.get("personality_model", {})
    if not isinstance(pm, dict) or not pm:
        return 0

    existing_candidates = list(twin_data.get("research_candidate_topics") or [])
    if len(existing_candidates) >= _MAX_RESEARCH_CANDIDATE_TOPICS:
        return 0

    gaps = twin_data.get("knowledge_gaps") or []
    gaps_summary = "\n".join(
        f"- {', '.join(g.get('topic_tags', []))} (asked {g.get('count', 0)}x)"
        for g in sorted(gaps, key=lambda g: int(g.get("count", 0)), reverse=True)[:10]
    ) or "None recorded."

    existing_summary = "\n".join(f"- {c.get('question_hint', '')}" for c in existing_candidates) or "None yet."

    # _context holds raw onboarding/deepen answer text, often long — excluded
    # to keep the prompt focused on the synthesized model itself.
    pm_for_prompt = {k: v for k, v in pm.items() if k != "_context"}
    title = twin_data.get("title", "")

    prompt_text = _RESEARCH_AGENT_PROMPT_TEMPLATE.format(
        name=twin_data.get("name", "this person"),
        title_suffix=f" ({title})" if title else "",
        personality_model_json=json.dumps(pm_for_prompt, indent=2)[:6000],
        knowledge_gaps_summary=gaps_summary,
        existing_candidates_summary=existing_summary,
    )

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": prompt_text}]}],
            inferenceConfig={"maxTokens": 800, "temperature": 0.9, "topP": 0.95},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        proposals = _extract_json_array(raw)
    except Exception as exc:
        print(f"[research-agent] Failed for twin {twin_id}: {exc}")
        return 0

    if not isinstance(proposals, list):
        return 0

    now = datetime.now().isoformat()
    added = 0
    for p in proposals:
        if not isinstance(p, dict) or not p.get("question_hint"):
            continue
        existing_candidates.append({
            "id": "RESEARCH:" + uuid.uuid4().hex[:12],
            "question_hint": str(p["question_hint"])[:400],
            "target_fields": (
                [f for f in (p.get("target_fields") or []) if isinstance(f, str)]
                or _GAP_TOPIC_DEFAULT_TARGET_FIELDS
            ),
            "rationale": str(p.get("rationale", ""))[:300],
            "added_at": now,
        })
        added += 1

    if added == 0:
        return 0

    if len(existing_candidates) > _MAX_RESEARCH_CANDIDATE_TOPICS:
        existing_candidates = existing_candidates[-_MAX_RESEARCH_CANDIDATE_TOPICS:]

    twin_data["research_candidate_topics"] = existing_candidates
    owner_id = twin_data.get("user_id")
    if owner_id:
        _save_twin(twin_id, owner_id, twin_data)
    else:
        _save_twin_flat(twin_id, twin_data)

    return added


def _list_all_twin_ids() -> List[str]:
    """Enumerate every twin record — user-created twins plus the migrated
    default twin, whichever storage backend is active. Built-in public
    personas (Gandhi et al.) are static files, not S3/local twin records, so
    they never appear here — nothing to research-agent for those anyway."""
    if USE_S3:
        ids = []
        paginator = s3_client.get_paginator("list_objects_v2")
        # Delimiter="/" returns only the flat twins/{twin_id}.json keys in
        # Contents, treating twins/{user_id}/ as a "folder" collapsed into
        # CommonPrefixes (ignored) — avoids double-counting each twin via its
        # per-user listing copy.
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=TWINS_S3_PREFIX, Delimiter="/"):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                if key.startswith(TWINS_S3_PREFIX) and key.endswith(".json"):
                    ids.append(key[len(TWINS_S3_PREFIX):-len(".json")])
        return ids
    twins_path = Path(TWINS_DIR)
    if not twins_path.exists():
        return []
    return [f.stem for f in twins_path.glob("*.json")]


async def run_research_agent_batch() -> dict:
    """Run the research agent for every twin. Invoked on a schedule (twice
    monthly — see terraform's aws_scheduler_schedule), not per-request, since
    this is exploratory analysis that doesn't need to be fresh turn-to-turn.
    Each twin is isolated in its own try/except so one bad twin record can't
    abort the batch."""
    twin_ids = _list_all_twin_ids()
    total_added = 0
    failures = 0
    for twin_id in twin_ids:
        try:
            total_added += await run_research_agent(twin_id)
        except Exception as exc:
            failures += 1
            print(f"[research-agent-batch] Failed for twin {twin_id}: {exc}")
    result = {"twins_processed": len(twin_ids), "candidates_added": total_added, "failures": failures}
    print(f"[research-agent-batch] {result}")
    return result


def _append_model_version(twin_data: dict, model_snapshot: dict, trigger: str) -> None:
    """Append a personality_model snapshot to the twin's version history,
    always retaining version 1 (the origin) even under cap eviction."""
    versions = list(twin_data.get("personality_model_versions") or [])
    next_version_num = (versions[-1]["version"] + 1) if versions else 1
    versions.append({
        "version": next_version_num,
        "model_snapshot": copy.deepcopy(model_snapshot),
        "created_at": datetime.now().isoformat(),
        "trigger": trigger,
    })
    if len(versions) > _MODEL_VERSIONS_CAP:
        origin = versions[0] if versions[0].get("version") == 1 else None
        rest = versions[1:] if origin else versions
        keep = _MODEL_VERSIONS_CAP - (1 if origin else 0)
        rest = rest[-keep:] if keep > 0 else []
        versions = ([origin] if origin else []) + rest
    twin_data["personality_model_versions"] = versions


def restore_personality_model_version(twin_data: dict, version: int) -> bool:
    """Restore twin_data's active personality_model to a previously recorded
    version. Snapshots the current (about-to-be-discarded) model first, then
    records the rollback itself as a new version entry. Returns False if the
    requested version doesn't exist."""
    versions = twin_data.get("personality_model_versions") or []
    match = next((v for v in versions if v.get("version") == version), None)
    if not match:
        return False
    current_model = twin_data.get("personality_model", {})
    _append_model_version(twin_data, current_model, trigger="pre-rollback")
    restored = copy.deepcopy(match["model_snapshot"])
    twin_data["personality_model"] = restored
    _append_model_version(twin_data, restored, trigger=f"rollback_to_v{version}")
    return True


class DeepenHistoryItem(BaseModel):
    role: str
    content: str


class DeepenRequest(BaseModel):
    history: List[DeepenHistoryItem] = Field(default_factory=list)
    # Accumulated raw answer text for whatever topic the previous response's
    # topic_id was about — echoed back so the endpoint can keep building on it
    # across turns until the critic judges it sufficient. Reset to "" by the
    # frontend whenever a response comes back with a different topic_id (or
    # topic_just_saved=true) signaling the topic changed.
    topic_answer_so_far: str = ""

    model_config = {"extra": "ignore"}


class DeepenResponse(BaseModel):
    message: str
    topic_id: Optional[str] = None
    topic_question_hint: Optional[str] = None
    topic_answer_so_far: str = ""
    topic_just_saved: bool = False
    topics_remaining_estimate: int = 0
    done: bool = False

    model_config = {"extra": "ignore"}

    @field_validator("done", "topic_just_saved", mode="before")
    @classmethod
    def _coerce_bool(cls, v: Any) -> Any:
        return v if isinstance(v, bool) else False


async def _patch_personality_model(twin_data: dict, ctx: dict, addressed_topic_ids: List[str]) -> None:
    """Targeted, additive update: ask Bedrock to return ONLY the fields
    relevant to the topics addressed this session, then merge just those
    fields into the existing model — everything else is left byte-for-byte
    unchanged, unlike the old full-model regeneration."""
    existing_model = twin_data.get("personality_model", {})
    target_fields = _target_fields_for_topics(addressed_topic_ids, twin_data)

    new_data_lines = []
    if ctx.get("pastDecisions"):
        new_data_lines.append(f"Past decisions: {ctx['pastDecisions'][-800:]}")
    if ctx.get("nonNegotiables"):
        new_data_lines.append(f"Non-negotiables: {ctx['nonNegotiables']}")
    if ctx.get("softPreferences"):
        new_data_lines.append(f"What they'd compromise on: {ctx['softPreferences']}")
    if ctx.get("mindChange"):
        new_data_lines.append(f"Changed their mind: {ctx['mindChange']}")
    if ctx.get("topicAnswer"):
        new_data_lines.append(f"New answer: {ctx['topicAnswer']}")
    new_data_text = "\n".join(new_data_lines) or "N/A"

    synthesis_prompt = f"""You are updating specific fields of an AI twin's personality model with new depth \
data. Do NOT regenerate the whole model — return ONLY the fields listed below, updated to incorporate the \
new data.

EXISTING MODEL (for context only — do not repeat fields not listed below):
{json.dumps(existing_model, indent=2)}

NEW DEPTH DATA:
{new_data_text}

FIELDS TO UPDATE (return ONLY these keys, each with its full updated value in the same shape as the \
existing model):
{", ".join(target_fields)}

Return ONLY a valid JSON object containing exactly these keys. No markdown, no extra text, no other fields."""

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            system=[{"text": synthesis_prompt}],
            messages=[{"role": "user", "content": [{"text": "Update the listed fields with the new depth data."}]}],
            inferenceConfig={"maxTokens": 800, "temperature": 0.5, "topP": 0.9},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        patch = _extract_json_object(raw)
        if isinstance(patch, dict):
            for field in target_fields:
                if field in patch and patch[field]:
                    existing_model[field] = patch[field]
    except Exception as exc:
        print(f"Deepen targeted patch failed (non-fatal): {exc}")
        # _context is already merged into existing_model by the caller, so
        # the raw depth data is saved even if this patch fails.

    existing_model["_context"] = ctx
    twin_data["personality_model"] = existing_model


async def _consolidate_personality_model(twin_data: dict, ctx: dict) -> None:
    """Periodic full re-synthesis pass (every _CONSOLIDATION_EVERY_N_PATCHES
    accepted targeted patches) to resolve any drift those incremental patches
    introduced between fields. This is the same full-regeneration prompt the
    old always-on path used — a version snapshot is taken by the caller
    before this runs, so a worse regeneration here is always recoverable."""
    existing_model = twin_data.get("personality_model", {})
    synthesis_prompt = f"""You are updating an AI twin's personality model with new depth data.

EXISTING MODEL:
{json.dumps(existing_model, indent=2)}

NEW DEPTH DATA:
Past decisions: {ctx.get("pastDecisions", "N/A")}
Non-negotiables (won't bend on): {ctx.get("nonNegotiables", "N/A")}
What they'd compromise on: {ctx.get("softPreferences", "N/A")}
Changed their mind: {ctx.get("mindChange", "N/A")}

Using the existing model as a base, return an improved JSON model that incorporates the new data and \
resolves any inconsistencies between fields that may have built up over incremental updates.
The new data should sharpen: decision_heuristics, blind_spots, what_they_avoid, decision_framework, personality_summary.
Preserve fields unaffected by this data: core_values, communication_traits, risk_profile, what_they_optimize_for.
Return ONLY valid JSON with the same structure as the existing model. No markdown, no extra text."""

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            system=[{"text": synthesis_prompt}],
            messages=[{"role": "user", "content": [{"text": "Update the personality model with the new depth data."}]}],
            inferenceConfig={"maxTokens": 1200, "temperature": 0.5, "topP": 0.9},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        updated_model = _extract_json_object(raw)
        if isinstance(updated_model, dict) and updated_model:
            updated_model["_context"] = ctx
            twin_data["personality_model"] = updated_model
    except Exception as exc:
        print(f"Deepen consolidation re-synthesis failed (non-fatal): {exc}")


async def _deepen_and_save(
    twin_id: str,
    user_id: str,
    twin_data: dict,
    new_fields: dict,
    addressed_topic_ids: List[str],
) -> None:
    """Merge new depth data into the twin's personality_model._context, then
    apply either a targeted field patch or (periodically) a full
    consolidation pass — see _patch_personality_model / _consolidate_personality_model.
    Snapshots a version before mutating, and decays any gap-ledger entries
    this session addressed."""
    existing_model = twin_data.get("personality_model", {})
    # context.py reads depth fields from personality_model["_context"], so persist there
    ctx = dict(existing_model.get("_context", {}))

    for key in ("pastDecisions", "nonNegotiables", "softPreferences", "mindChange", "topicAnswer"):
        if new_fields.get(key):
            if key == "pastDecisions" and ctx.get(key):
                ctx[key] = ctx[key].strip() + "\n\n" + new_fields[key].strip()
            else:
                ctx[key] = new_fields[key]

    # Write the updated _context back into personality_model so prompt building picks it up
    existing_model["_context"] = ctx
    twin_data["personality_model"] = existing_model
    twin_data["sources"] = merge_sources(
        twin_data.get("sources", []),
        build_deepen_sources(new_fields),
    )

    _decay_addressed_gaps(twin_data, addressed_topic_ids)
    _consume_research_candidates(twin_data, addressed_topic_ids)

    # Snapshot before mutating — on a twin's very first deepen update, this
    # also seeds version 1 (the origin) with the pre-deepen model.
    _append_model_version(twin_data, existing_model, trigger="pre-deepen-update")

    patch_count = int(twin_data.get("deepen_patch_count", 0)) + 1
    twin_data["deepen_patch_count"] = patch_count

    if patch_count % _CONSOLIDATION_EVERY_N_PATCHES == 0:
        await _consolidate_personality_model(twin_data, ctx)
    else:
        await _patch_personality_model(twin_data, ctx, addressed_topic_ids)

    twin_data["deepen_completed_at"] = datetime.now().isoformat()
    try:
        _save_twin(twin_id, user_id, twin_data)
    except Exception as exc:
        print(f"Deepen save failed: {exc}")
        raise


@app.post("/twin/{twin_id}/deepen/message", response_model=DeepenResponse)
async def deepen_message(
    twin_id: str,
    request: DeepenRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Run one turn of the deepen interview for a twin the caller owns.

    Single-topic focused: instead of offering a batch of topics and waiting
    until an entire session is "done" to save anything, each call focuses on
    exactly one topic (_select_deepen_topics()[0]) and keeps following up on
    it — via a separate critic judgment (_judge_topic_sufficiency) — until the
    accumulated answer is judged concrete enough. The moment that happens, the
    field is patched and saved immediately (_deepen_and_save), before moving
    to the next topic. This means navigating away mid-session only loses
    progress on whatever topic is currently in flight, not everything covered
    so far — and because saved topics stop being thin/gapped, they're
    naturally never re-offered on a later call without needing any separate
    "already covered" ledger.

    Topic selection is gap-driven and covers the whole personality model, not
    a fixed list (see openspec/changes/generative-deepen-interview and the
    broadened _THIN_FIELD_DEEPEN_TOPICS). A session ending (done=true) never
    permanently closes the twin to future deepening — recomputed fresh every
    call, so a twin can always be deepened again once new gaps, thin fields,
    or research-agent candidate topics exist.
    """
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    pm = twin_data.get("personality_model", {})
    ctx_data = pm.get("_context", {}) if isinstance(pm, dict) else {}

    remaining = _select_deepen_topics(twin_data, [], limit=10)

    # If nothing is selectable right now, say so honestly without writing any
    # permanent "fully deepened" state — a future call (once new gaps or thin
    # fields exist, or the twice-monthly research agent adds candidates) will
    # find something again.
    if not remaining:
        return {
            "message": (
                f"Nothing new to deepen for {twin_data.get('name', 'this twin')} right now — "
                "check back after a few more conversations, or once there's new context to draw on."
            ),
            "topic_id": None,
            "topic_answer_so_far": "",
            "topic_just_saved": False,
            "topics_remaining_estimate": 0,
            "done": True,
        }

    current_topic = remaining[0]
    topics_remaining_estimate = len(remaining)

    ctx_lines = []
    if twin_data.get("name"):
        ctx_lines.append(f"Name: {twin_data['name']}")
    if twin_data.get("title"):
        ctx_lines.append(f"Title: {twin_data['title']}")
    if isinstance(pm, dict):
        if pm.get("personality_summary"):
            ctx_lines.append(f"Personality: {pm['personality_summary']}")
        if pm.get("decision_framework"):
            ctx_lines.append(f"Decision framework: {pm['decision_framework']}")
    if isinstance(ctx_data, dict):
        if ctx_data.get("pastDecisions"):
            ctx_lines.append(f"Past decisions (already captured): {ctx_data['pastDecisions'][:300]}")
        if ctx_data.get("nonNegotiables"):
            ctx_lines.append(f"Non-negotiables (already captured): {ctx_data['nonNegotiables'][:200]}")
        if ctx_data.get("mindChange"):
            ctx_lines.append(f"Mind change (already captured): {ctx_data['mindChange'][:200]}")
    existing_context = "\n".join(ctx_lines) if ctx_lines else "No existing context."

    system_prompt = _DEEPEN_TOPIC_SYSTEM_TEMPLATE.format(
        topic_hint=current_topic["question_hint"],
        existing_context=existing_context,
    )

    messages: List[Dict[str, Any]] = []
    if not request.history:
        messages = [{"role": "user", "content": [{"text": "hi, let's start"}]}]
    else:
        for item in request.history[-30:]:
            if item.role not in ("user", "assistant"):
                raise HTTPException(status_code=400, detail=f"Invalid role: {item.role!r}")
            messages.append({"role": item.role, "content": [{"text": item.content}]})
        if messages and messages[0]["role"] != "user":
            messages.insert(0, {"role": "user", "content": [{"text": "hi, let's start"}]})

    fallback_message = "Got it — let me keep going. Could you tell me more?"
    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            system=[{"text": system_prompt}],
            messages=messages,
            inferenceConfig={"maxTokens": 400, "temperature": 0.9, "topP": 0.95},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        data = _extract_json_object(raw, required_key="message")

        if not isinstance(data, dict) or "message" not in data:
            raise ValueError("Invalid deepen JSON structure")

        reply_message = str(data.get("message") or fallback_message)
        answer_delta = str(data.get("answer_delta") or "").strip()

    except (ValueError, json.JSONDecodeError) as exc:
        print(f"Deepen JSON parse error: {exc!r}")
        reply_message = raw if "raw" in locals() and raw.strip() and not raw.strip().startswith("{") else fallback_message
        answer_delta = ""
    except Exception as exc:
        print(f"Unexpected error in /twin/{{twin_id}}/deepen/message: {exc}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred. Please try again.")

    accumulated_answer = (
        (request.topic_answer_so_far.strip() + "\n" + answer_delta).strip()
        if answer_delta else request.topic_answer_so_far.strip()
    )

    is_sufficient, _reason = await _judge_topic_sufficiency(current_topic["question_hint"], accumulated_answer)

    if not is_sufficient:
        return {
            "message": reply_message,
            "topic_id": current_topic["id"],
            "topic_question_hint": current_topic["question_hint"],
            "topic_answer_so_far": accumulated_answer,
            "topic_just_saved": False,
            "topics_remaining_estimate": topics_remaining_estimate,
            "done": False,
        }

    # Sufficient — save this single topic's field(s) immediately rather than
    # waiting for the whole session to end. This is the fix for the data-loss
    # bug: whatever's saved here survives a navigate-away, refresh, or crash.
    evergreen_meta = _EVERGREEN_DEEPEN_TOPICS.get(current_topic["id"])
    field_key = evergreen_meta["field_update_key"] if evergreen_meta else "topicAnswer"
    new_fields = {field_key: accumulated_answer}
    await _deepen_and_save(twin_id, user_id, twin_data, new_fields, [current_topic["id"]])

    next_remaining = _select_deepen_topics(twin_data, [], limit=10)
    if next_remaining:
        next_topic = next_remaining[0]
        transition_message = f"Got it — that's genuinely useful.\n\n{next_topic['question_hint']}"
        return {
            "message": transition_message,
            "topic_id": next_topic["id"],
            "topic_question_hint": next_topic["question_hint"],
            "topic_answer_so_far": "",
            "topic_just_saved": True,
            "topics_remaining_estimate": len(next_remaining),
            "done": False,
        }

    return {
        "message": f"Got it — that's genuinely useful. That's everything worth covering for {twin_data.get('name', 'this twin')} right now — nice work.",
        "topic_id": None,
        "topic_answer_so_far": "",
        "topic_just_saved": True,
        "topics_remaining_estimate": 0,
        "done": True,
    }


class PersonalityRollbackRequest(BaseModel):
    version: int


@app.patch("/twin/{twin_id}/personality/rollback")
async def rollback_personality_model(
    twin_id: str,
    request: PersonalityRollbackRequest,
    user_id: str = Depends(get_current_user_id),
):
    """Restore a twin's personality_model to a previously recorded version.

    Targeted deepen patches and periodic consolidation passes can degrade a
    field instead of improving it; this lets an owner undo a specific update
    without losing everything since — see personality_model_versions."""
    twin_data = load_twin(twin_id)
    if not twin_data or twin_data.get("user_id") != user_id:
        raise HTTPException(status_code=404, detail="Twin not found")

    if not restore_personality_model_version(twin_data, request.version):
        raise HTTPException(status_code=404, detail=f"Version {request.version} not found")

    _save_twin(twin_id, user_id, twin_data)
    return {
        "status": "ok",
        "active_version": twin_data["personality_model_versions"][-1]["version"],
        "versions_available": [v["version"] for v in twin_data["personality_model_versions"]],
    }


# ── Resume Builder ─────────────────────────────────────────────────────────────

_RESUME_SYSTEM_TEMPLATE = """\
You are a focused assistant collecting data to build someone's professional resume.
Cover the remaining topics through natural conversation — one topic per turn.

TOPICS (in order, skip already-covered ones):
1. TECH_STACK       → languages, frameworks, tools, cloud platforms, certifications
2. EDUCATION        → degrees, institutions, fields of study, graduation years
3. CAREER_HISTORY   → companies, job titles, start/end dates, key responsibilities
4. ACCOMPLISHMENTS  → 2-4 specific wins with metrics (e.g. "reduced latency by 40%")
5. TARGET_ROLE      → what kind of role they are targeting next

RULES:
- One topic per turn. 2-3 sentences max. Be direct and practical.
- Every message MUST end with a question, except the final closing when done is true.
- If an answer is vague (e.g. "I've done a lot of backend work"), ask for one specific \
example or metric, then accept whatever they give.
- After a complete answer, acknowledge in one short phrase and ask the next topic.
- Never use form-speak ("Topic 3 of 5", "Moving on to", "Next up").
- When all topics are covered, close with one sentence and set done to true.

CURRENT STATE:
Topics remaining: {topics_remaining}
Data collected so far:
{fields_json}
{linkedin_section}

RETURN ONLY valid JSON — no markdown, no text outside the JSON:
{{
  "message": "your question or closing as natural prose",
  "field_updates": {{
    "tech_stack": "comma-separated tools/languages — omit if not in this turn",
    "education": "degree, institution, year — omit if not in this turn",
    "career_history": "bullet list of roles: Company (dates): responsibilities — omit if not in this turn",
    "accomplishments": "bullet list of wins with metrics — omit if not in this turn",
    "target_role": "desired role title — omit if not in this turn"
  }},
  "topics_covered": ["TECH_STACK"],
  "done": false
}}
"""

_ALL_RESUME_TOPICS = ["TECH_STACK", "EDUCATION", "CAREER_HISTORY", "ACCOMPLISHMENTS", "TARGET_ROLE"]

_RESUME_TOOLS = [
    {
        "toolSpec": {
            "name": "set_contact_info",
            "description": "Set the candidate's contact information for the resume header.",
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string"},
                        "email": {"type": "string"},
                        "phone": {"type": "string"},
                        "location": {"type": "string"},
                        "linkedin_url": {"type": "string"},
                    },
                    "required": ["name"],
                },
            },
        }
    },
    {
        "toolSpec": {
            "name": "set_summary",
            "description": (
                "Write a 2-4 sentence professional summary. "
                "If a target role and job description are provided, tailor it to that role."
            ),
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "text": {"type": "string"},
                    },
                    "required": ["text"],
                },
            },
        }
    },
    {
        "toolSpec": {
            "name": "add_experience",
            "description": "Add one work experience entry. Call once per role, most recent first.",
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "company": {"type": "string"},
                        "title": {"type": "string"},
                        "start_date": {"type": "string"},
                        "end_date": {"type": "string", "description": "Use 'Present' if current role"},
                        "bullets": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": (
                                "2-4 achievement-oriented bullet points. "
                                "Start each with a strong action verb. Include metrics where available."
                            ),
                        },
                    },
                    "required": ["company", "title", "start_date", "end_date", "bullets"],
                },
            },
        }
    },
    {
        "toolSpec": {
            "name": "add_education",
            "description": "Add one education entry. Call once per degree.",
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "institution": {"type": "string"},
                        "degree": {"type": "string"},
                        "field": {"type": "string"},
                        "graduation_year": {"type": "string"},
                        "gpa": {"type": "string", "description": "Include only if 3.5 or above"},
                    },
                    "required": ["institution", "degree", "graduation_year"],
                },
            },
        }
    },
    {
        "toolSpec": {
            "name": "set_skills",
            "description": "Set the skills section grouped by category (e.g. Languages, Frameworks, Tools, Cloud).",
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "categories": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "name": {"type": "string"},
                                    "skills": {"type": "array", "items": {"type": "string"}},
                                },
                                "required": ["name", "skills"],
                            },
                        },
                    },
                    "required": ["categories"],
                },
            },
        }
    },
    {
        "toolSpec": {
            "name": "finalize_resume",
            "description": "Call this once all sections have been populated. Signals the backend to format and return the document.",
            "inputSchema": {
                "json": {
                    "type": "object",
                    "properties": {
                        "ready": {"type": "boolean"},
                    },
                    "required": ["ready"],
                },
            },
        }
    },
]


class ResumeHistoryItem(BaseModel):
    role: str
    content: str


class ResumeFieldUpdates(BaseModel):
    tech_stack: Optional[str] = None
    education: Optional[str] = None
    career_history: Optional[str] = None
    accomplishments: Optional[str] = None
    target_role: Optional[str] = None

    model_config = {"extra": "ignore"}


class ResumeRequest(BaseModel):
    history: List[ResumeHistoryItem] = Field(default_factory=list)
    topics_covered: List[str] = Field(default_factory=list)
    fields_collected: Optional[Dict[str, Any]] = None
    linkedin_parsed: Optional[Dict[str, Any]] = None


class ResumeResponse(BaseModel):
    message: str
    field_updates: ResumeFieldUpdates = Field(default_factory=ResumeFieldUpdates)
    topics_covered: List[str] = Field(default_factory=list)
    done: bool = False

    model_config = {"extra": "ignore"}

    @field_validator("field_updates", mode="before")
    @classmethod
    def _coerce_field_updates(cls, v: Any) -> Any:
        if not isinstance(v, dict):
            return {}
        return v

    @field_validator("topics_covered", mode="before")
    @classmethod
    def _coerce_topics_covered(cls, v: Any) -> Any:
        if not isinstance(v, list):
            return []
        return [item for item in v if isinstance(item, str)]

    @field_validator("done", mode="before")
    @classmethod
    def _coerce_done(cls, v: Any) -> Any:
        if isinstance(v, bool):
            return v
        return False


class ResumeGenerateRequest(BaseModel):
    fields_collected: Dict[str, Any] = Field(default_factory=dict)


def _build_resume_docx(resume_data: dict) -> bytes:
    """Format accumulated tool-call results into a Word document."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    contact = resume_data.get("contact", {})

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(contact.get("name", ""))
    run.bold = True
    run.font.size = Pt(18)

    # Contact line
    parts = [v for k in ("email", "phone", "location", "linkedin_url") if (v := contact.get(k))]
    if parts:
        cp = doc.add_paragraph(" | ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        for r in cp.runs:
            r.font.size = Pt(10)

    def _section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Summary
    if resume_data.get("summary"):
        _section_header("Professional Summary")
        p = doc.add_paragraph(resume_data["summary"])
        p.paragraph_format.space_after = Pt(4)

    # Experience
    if resume_data.get("experience"):
        _section_header("Experience")
        for exp in resume_data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(exp.get("company", ""))
            r.bold = True
            r.font.size = Pt(11)
            p.add_run(f"  {exp.get('start_date', '')} – {exp.get('end_date', '')}")

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(exp.get("title", ""))
            r2.italic = True

            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.add_run(bullet)

    # Education
    if resume_data.get("education"):
        _section_header("Education")
        for edu in resume_data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(edu.get("institution", ""))
            r.bold = True
            r.font.size = Pt(11)
            p.add_run(f"  {edu.get('graduation_year', '')}")

            degree_parts = [s for s in (edu.get("degree"), edu.get("field")) if s]
            if edu.get("gpa"):
                degree_parts.append(f"GPA: {edu['gpa']}")
            if degree_parts:
                p2 = doc.add_paragraph(", ".join(degree_parts))
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(4)

    # Skills
    if resume_data.get("skills"):
        _section_header("Skills")
        for cat in resume_data["skills"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f"{cat.get('name', '')}: ")
            r.bold = True
            p.add_run(", ".join(cat.get("skills", [])))

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.post("/resume/message")
async def resume_message(
    request: ResumeRequest,
    _user_id: str = Depends(get_current_user_id),
):
    """One turn of the resume data-collection interview."""
    covered = list(request.topics_covered)

    # Auto-mark topics that are already satisfied by pre-loaded twin/LinkedIn data
    fields = request.fields_collected or {}
    lp = request.linkedin_parsed or {}
    if lp.get("skills") or fields.get("tech_stack"):
        if "TECH_STACK" not in covered:
            covered.append("TECH_STACK")
    if fields.get("career_history") or lp.get("experience"):
        if "CAREER_HISTORY" not in covered:
            covered.append("CAREER_HISTORY")
    if fields.get("accomplishments") or lp.get("achievements"):
        if "ACCOMPLISHMENTS" not in covered:
            covered.append("ACCOMPLISHMENTS")

    remaining = [t for t in _ALL_RESUME_TOPICS if t not in covered]

    linkedin_section = ""
    if lp:
        lines = []
        if lp.get("name"):       lines.append(f"Name: {lp['name']}")
        if lp.get("title"):      lines.append(f"Title: {lp['title']}")
        if lp.get("skills"):     lines.append(f"Skills: {str(lp['skills'])[:300]}")
        if lp.get("experience"): lines.append(f"Experience: {str(lp['experience'])[:400]}")
        if lines:
            linkedin_section = "\nLinkedIn data (pre-loaded — do not re-ask):\n" + "\n".join(lines)

    system_prompt = _RESUME_SYSTEM_TEMPLATE.format(
        topics_remaining=", ".join(remaining) if remaining else "ALL COVERED",
        fields_json=json.dumps(fields, indent=2),
        linkedin_section=linkedin_section,
    )

    messages: List[Dict[str, Any]] = []
    if not request.history:
        messages = [{"role": "user", "content": [{"text": "hi, let's start"}]}]
    else:
        for item in request.history[-50:]:
            if item.role not in ("user", "assistant"):
                raise HTTPException(status_code=400, detail=f"Invalid role: {item.role!r}")
            messages.append({"role": item.role, "content": [{"text": item.content}]})
        if messages and messages[0]["role"] != "user":
            messages.insert(0, {"role": "user", "content": [{"text": "hi, let's start"}]})

    raw: Optional[str] = None
    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            system=[{"text": system_prompt}],
            messages=messages,
            inferenceConfig={"maxTokens": 600, "temperature": 0.8, "topP": 0.95},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        data = _extract_json_object(raw, required_key="message")

        if not isinstance(data, dict) or "message" not in data:
            raise ValueError("Invalid resume interview JSON structure")

        validated = ResumeResponse.model_validate(data)

        request_topics = set(covered)
        model_topics = set(validated.topics_covered or [])
        merged = request_topics | model_topics
        if merged:
            validated.topics_covered = [t for t in _ALL_RESUME_TOPICS if t in merged]

        result = validated.model_dump(exclude_none=True)
        if set(_ALL_RESUME_TOPICS).issubset(merged) and not result.get("done"):
            result["done"] = True

        return result

    except (ValueError, json.JSONDecodeError) as exc:
        print(f"Resume interview JSON parse error: {exc!r}")
        done_msg = "Perfect — I have enough to generate your resume now."
        cont_msg = "Got it — let me keep going. Could you tell me a bit more?"

        covered_set = set(covered)
        fallback_done = set(_ALL_RESUME_TOPICS).issubset(covered_set)
        fallback_message = raw if raw is not None else cont_msg
        merged_fallback_topics = set(covered_set)

        if raw is not None:
            try:
                parsed = _extract_json_object(raw)
                if isinstance(parsed, dict):
                    if parsed.get("done") is True:
                        fallback_done = True
                    parsed_topics = parsed.get("topics_covered")
                    if isinstance(parsed_topics, list):
                        merged_fallback_topics |= {t for t in parsed_topics if t in _ALL_RESUME_TOPICS}
                    parsed_message = parsed.get("message")
                    if isinstance(parsed_message, str) and parsed_message.strip():
                        fallback_message = parsed_message.strip()
            except (ValueError, json.JSONDecodeError) as parse_exc:
                print(f"Resume interview fallback JSON parse error: {parse_exc!r}")

            if raw.strip().startswith("{") or raw.strip().startswith("```"):
                if fallback_message == raw:
                    fallback_message = cont_msg
            else:
                last_brace = raw.rfind('{')
                if last_brace > len(raw) // 2:
                    try:
                        fragment = json.loads(raw[last_brace:])
                        if isinstance(fragment, dict) and "done" in fragment:
                            if fragment.get("done") is True:
                                fallback_done = True
                            fallback_message = raw[:last_brace].strip()
                    except json.JSONDecodeError:
                        pass
                if not fallback_message:
                    fallback_message = done_msg if fallback_done else cont_msg

        if set(_ALL_RESUME_TOPICS).issubset(merged_fallback_topics):
            fallback_done = True
        fallback_topics = (
            list(_ALL_RESUME_TOPICS)
            if fallback_done
            else [t for t in _ALL_RESUME_TOPICS if t in merged_fallback_topics]
        )
        return {
            "message": fallback_message,
            "field_updates": {},
            "topics_covered": fallback_topics,
            "done": fallback_done,
        }
    except Exception as exc:
        print(f"Unexpected error in /resume/message: {exc}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred. Please try again.")


@app.post("/resume/parse-jd")
async def resume_parse_jd(
    file: UploadFile = File(...),
    _user_id: str = Depends(get_current_user_id),
):
    """Extract structured data from a job description PDF."""
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    try:
        contents = await file.read()
        reader = PdfReader(io.BytesIO(contents))
        text = "".join(p.extract_text() or "" for p in reader.pages)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {e}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from PDF")

    extract_prompt = f"""Extract structured data from this job description and return ONLY valid JSON:
{{
  "role": "job title",
  "company": "company name or empty string",
  "required_skills": "comma-separated list of required skills/technologies",
  "preferred_skills": "comma-separated nice-to-haves",
  "responsibilities": "bullet summary of key responsibilities",
  "raw_text": "first 800 chars of the job description verbatim"
}}

Job description:
{text[:5000]}

Return only the JSON, no other text."""

    try:
        response = await asyncio.to_thread(
            bedrock_client.converse,
            modelId=BEDROCK_MODEL_ID,
            messages=[{"role": "user", "content": [{"text": extract_prompt}]}],
            inferenceConfig={"maxTokens": 800, "temperature": 0.1},
        )
        raw = response["output"]["message"]["content"][0]["text"]
        return _extract_json_object(raw)
    except Exception as e:
        print(f"Error in /resume/parse-jd: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse job description")


@app.post("/resume/generate")
async def resume_generate(
    request: ResumeGenerateRequest,
    _user_id: str = Depends(get_current_user_id),
):
    """Run the resume-builder agent loop, then return a formatted .docx file."""
    fields = request.fields_collected

    context_lines = []
    for key, val in fields.items():
        if val and key != "job_description":
            context_lines.append(f"{key.replace('_', ' ').title()}: {val}")
    context = "\n".join(context_lines)

    jd_section = ""
    if fields.get("job_description"):
        jd_section = f"\n\nTarget Job Description:\n{fields['job_description']}"

    system_prompt = (
        "You are a professional resume writer. Use the provided tools to build a polished, "
        "ATS-friendly resume from the candidate information supplied.\n\n"
        "Call the tools in this order:\n"
        "1. set_contact_info\n"
        "2. set_summary (tailor to the target role if a job description is provided)\n"
        "3. add_experience — once per role, most recent first; write strong action-verb bullets with metrics\n"
        "4. add_education — once per degree\n"
        "5. set_skills — group by category (Languages, Frameworks, Tools, Cloud, etc.)\n"
        "6. finalize_resume\n\n"
        "Use only information that is explicitly provided. Do not invent details, dates, or metrics."
    )

    user_message = f"Here is the candidate's information:\n\n{context}{jd_section}\n\nBuild their resume now."

    messages: List[Dict[str, Any]] = [
        {"role": "user", "content": [{"text": user_message}]}
    ]

    resume_data: Dict[str, Any] = {
        "contact": {},
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
    }

    _MAX_AGENT_ITERATIONS = 20

    try:
        for _ in range(_MAX_AGENT_ITERATIONS):
            response = await asyncio.to_thread(
                bedrock_client.converse,
                modelId=BEDROCK_MODEL_ID,
                system=[{"text": system_prompt}],
                messages=messages,
                toolConfig={
                    "tools": _RESUME_TOOLS,
                    "toolChoice": {"auto": {}},
                },
                inferenceConfig={"maxTokens": 4096, "temperature": 0.2},
            )

            stop_reason = response.get("stopReason", "end_turn")
            response_message = response["output"]["message"]
            # Bedrock sometimes returns empty text blocks alongside toolUse blocks;
            # filter them out to avoid ValidationException on the next converse call.
            if response_message.get("content"):
                response_message["content"] = [
                    b for b in response_message["content"]
                    if not (isinstance(b.get("text"), str) and not b["text"].strip())
                ]
            messages.append(response_message)

            if stop_reason != "tool_use":
                break

            tool_results = []
            finalized = False

            for block in response_message.get("content", []):
                tool_use = block.get("toolUse")
                if not tool_use:
                    continue

                name = tool_use["name"]
                inp = tool_use["input"]
                tool_use_id = tool_use["toolUseId"]

                if name == "set_contact_info":
                    resume_data["contact"] = inp
                elif name == "set_summary":
                    resume_data["summary"] = inp.get("text", "")
                elif name == "add_experience":
                    resume_data["experience"].append(inp)
                elif name == "add_education":
                    resume_data["education"].append(inp)
                elif name == "set_skills":
                    resume_data["skills"] = inp.get("categories", [])
                elif name == "finalize_resume":
                    finalized = True

                tool_results.append({
                    "toolResult": {
                        "toolUseId": tool_use_id,
                        "content": [{"text": "OK"}],
                        "status": "success",
                    }
                })

            if tool_results:
                messages.append({"role": "user", "content": tool_results})

            if finalized:
                break

    except Exception as exc:
        print(f"Error in resume agent loop: {exc}")
        raise HTTPException(status_code=500, detail="Failed to generate resume")

    if not resume_data["contact"] and not resume_data["experience"]:
        raise HTTPException(status_code=500, detail="Resume agent did not produce usable output")

    try:
        doc_bytes = await asyncio.to_thread(_build_resume_docx, resume_data)
    except Exception as exc:
        print(f"Error building resume docx: {exc}")
        raise HTTPException(status_code=500, detail="Failed to format resume document")

    raw_candidate_name = str(resume_data["contact"].get("name") or "resume")
    safe_candidate_name = re.sub(r"[^A-Za-z0-9_-]+", "_", raw_candidate_name).strip("._-")
    safe_candidate_name = re.sub(r"_+", "_", safe_candidate_name)
    if not safe_candidate_name:
        safe_candidate_name = "resume"
    filename = f"{safe_candidate_name}_resume.docx"
    filename_star = f"{quote(safe_candidate_name, safe='')}_resume.docx"

    return StreamingResponse(
        io.BytesIO(doc_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\"; filename*=UTF-8''{filename_star}"},
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)  # nosec B104 — local dev only, not used in Lambda
